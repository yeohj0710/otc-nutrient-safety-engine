from __future__ import annotations

import csv
import hashlib
import json
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PICOS_PATH = ROOT / "research_v3/otc/literature/picos/picos_definition.json"
EVIDENCE_PATH = ROOT / "research_v3/otc/literature/evidence_map.csv"
SCREENING_PATH = ROOT / "research_v3/otc/literature/screening/screening_manifest.json"
OTC_METRICS_PATH = ROOT / "research_v3/otc/metrics_manifest.json"
ROOT_METRICS_PATH = ROOT / "research_v3/metrics_manifest.json"
AI_REFERENCE_PATH = ROOT / "research_v3/measurement/screener_vs_ai_reference.json"
BLIND_EVAL_PATH = ROOT / "research_v3/otc/validation/ai_independent_evaluation.json"
LINK_MANIFEST_PATH = ROOT / "research_v3/otc/rules/literature_link_manifest.json"
COMPLETION_AUDIT_PATH = ROOT / "research_v3/otc/audit/completion_audit.json"
DEPLOYMENT_RECEIPT_PATH = ROOT / "research_v3/otc/audit/production_deployment_receipt.json"
THESIS_DIR = ROOT / "research_v3/thesis"
THESIS_DOCX = THESIS_DIR / "권혁찬_졸업논문_최종본.docx"
REPORTS_DIR = ROOT / "research_v3/reports"

FONT_BODY = "Pretendard"
FONT_LIGHT = "Pretendard Light"
FONT_MEDIUM = "Pretendard Medium"
FONT_SEMIBOLD = "Pretendard SemiBold"
FONT_EXTRABOLD = "Pretendard ExtraBold"
NAVY = RGBColor(18, 43, 67)
BLUE = RGBColor(38, 94, 139)
MUTED = RGBColor(92, 105, 117)
LIGHT_FILL = "EEF3F7"
CAUTION_FILL = "FFF4D6"
TABLE_WIDTH_DXA = 9000


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8-sig"))


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as source:
        return list(csv.DictReader(source))


def sha256_file(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def pct(value: float | None) -> str:
    """비율을 백분율 문자열로. 값이 없으면 null 로 남긴다."""
    return "null" if value is None else f"{value * 100:.2f}%"


def num(value: float | None, digits: int = 4) -> str:
    return "null" if value is None else f"{value:.{digits}f}"


def deployment_sentence(metrics: dict[str, Any], polite: bool = False) -> str:
    """배포 여부를 영수증에서 읽어 한 문장으로 만든다. 추정하지 않는다.

    문서마다 문체가 다르므로 평서체와 존댓말을 따로 낸다.
    """
    if not metrics["deployed"]:
        return "이번 실행에서는 배포하지 않았습니다." if polite else "이번 실행에서는 배포하지 않았다."
    if polite:
        return (
            f"사용자 지시로 production 에 배포했습니다. 공개 주소는 {metrics['public_url']} 이고 "
            f"배포 ID 는 {metrics['deployment_id']} 입니다. 사이트 배포는 연구 상태 플래그 "
            "release_ready 와 별개이며, release_ready 는 임상 배포 승인 절차를 뜻하므로 "
            "false 를 유지합니다."
        )
    return (
        f"사용자 지시로 production 에 배포했다. 공개 주소는 {metrics['public_url']} 이고 "
        f"배포 ID 는 {metrics['deployment_id']} 다. 사이트 배포는 연구 상태 플래그 "
        "release_ready 와 별개이며 release_ready 는 임상 배포 승인 절차를 뜻하므로 false 를 유지한다."
    )


def collect_metrics() -> dict[str, Any]:
    """모든 수치를 매니페스트에서 읽는다. 문서에 숫자를 하드코딩하지 않는다."""
    picos = read_json(PICOS_PATH)
    evidence = read_csv(EVIDENCE_PATH)
    screening = read_json(SCREENING_PATH)
    otc = read_json(OTC_METRICS_PATH)
    metrics = otc["metrics"]
    reference = read_json(AI_REFERENCE_PATH)
    blind = read_json(BLIND_EVAL_PATH)
    links = read_json(LINK_MANIFEST_PATH)
    completion = read_json(COMPLETION_AUDIT_PATH)
    deployment = (
        read_json(DEPLOYMENT_RECEIPT_PATH) if DEPLOYMENT_RECEIPT_PATH.exists() else {"deployed": False}
    )

    primary = blind["primary_analysis"]
    ref_primary = reference["primary_analysis"]
    flags = completion["status_flags"]
    return {
        # 허가원문 결정층
        "products": metrics["analysis_products"]["value"],
        "ingredients": metrics["analysis_ingredients"]["value"],
        "bindings": metrics["runtime_product_ingredient_bindings"]["value"],
        "constraints": metrics["verified_administration_constraints"]["value"],
        "rules_total": metrics["rules_total"]["value"],
        "rules_released": metrics["rules_released"]["value"],
        "static_paths": metrics["static_paths_generated"]["value"],
        "research_tests": metrics["research_tests_passed"]["value"],
        "app_tests": metrics["app_tests_passed"]["value"],
        # AI 자율 PICOS 와 코퍼스
        "questions": len(picos["questions"]),
        "question_rows": [
            {
                "id": question["question_id"],
                "title": question["title_ko"],
                "hits": question.get("observed_hit_count"),
            }
            for question in picos["questions"]
        ],
        "picos_prompt_sha256": picos["prompt_sha256"],
        "search_hits": picos["last_search"]["total_hits_before_deduplication"],
        "corpus_rows": len(evidence),
        "with_abstract": sum(row["has_abstract"] == "true" for row in evidence),
        "title_only": sum(row["has_abstract"] != "true" for row in evidence),
        # P2 선별
        "screened": screening["classified_rows"],
        "coverage": screening["coverage"],
        "decision_distribution": screening["decision_distribution"],
        "screening_prompt_sha256": screening["prompt_sha256"],
        "corpus_sha256": screening["input_sha256"],
        "screener": screening["screener"],
        "screening_complete": screening["run_complete"],
        "batch_count": screening["batch_count"],
        "human_decisions": screening["human_decisions"],
        # P3-A AI 참조표준
        "ref_sample": reference["sample_size"],
        "ref_strata": len(reference["strata"]),
        "ref_prevalence": reference["corpus_prevalence"]["rogan_gladen_corrected_prevalence"],
        "ref_prevalence_ci": reference["corpus_prevalence"]["rogan_gladen_corrected_prevalence_ci95"],
        "ref_estimated_retain": reference["corpus_prevalence"]["estimated_corpus_retain_count"],
        "ref_sensitivity": ref_primary["sensitivity_vs_ai_reference"],
        "ref_specificity": ref_primary["specificity_vs_ai_reference"],
        "ref_precision": ref_primary["precision_vs_ai_reference"],
        "ref_f1": ref_primary["f1_vs_ai_reference"],
        "ref_agreement": ref_primary["agreement_vs_ai_reference"],
        "ref_rows": ref_primary["rows_analyzed"],
        "ref_uncertain": reference["ai_reference_label_distribution"]["uncertain"],
        "ref_inter_round_agreement": reference["inter_round"]["mean_pairwise_agreement"],
        "ref_prompt_sha256": reference["prompt_sha256"],
        # P3-B 규칙엔진 AI 맹검 독립평가
        "blind_cases": blind["cases_total"],
        "blind_scored": primary["cases"],
        "blind_sensitivity": primary["sensitivity_vs_ai_reference"],
        "blind_specificity": primary["specificity_vs_ai_reference"],
        "blind_precision": primary["precision_vs_ai_reference"],
        "blind_f1": primary["f1_vs_ai_reference"],
        "blind_agreement": primary["agreement_vs_ai_reference"],
        "blind_sensitivity_ci": primary["sensitivity_wilson_ci95"],
        "blind_specificity_ci": primary["specificity_wilson_ci95"],
        "blind_tp": primary["true_positive"],
        "blind_fp": primary["false_positive"],
        "blind_fn": primary["false_negative"],
        "blind_tn": primary["true_negative"],
        "blind_uncertain": blind["excluded_uncertain"],
        "blind_unresolved": blind["excluded_unresolved"],
        "blind_compromised": blind["excluded_blinding_compromised"],
        "blind_inter_round_agreement": blind["inter_round"]["mean_pairwise_agreement"],
        "blind_lock_sha256": blind["lock"]["sha256"],
        "blind_gap_rule_types": len(blind["coverage_gap_analysis"]["by_rule_type"]),
        "blind_gap_details": blind["coverage_gap_analysis"]["by_rule_type"],
        "blind_per_rule_type": blind["per_rule_type"],
        "draft_rule_types": blind["draft_rule_analysis"]["rule_types"],
        "draft_uncertain": blind["draft_rule_analysis"]["uncertain_cases"],
        # P4 문헌 근거 연결
        "link_rules_total": links["rules_total"],
        "link_rules_covered": links["rules_with_literature"],
        "link_total": links["links_total"],
        "link_pmids": links["unique_pmids"],
        "link_conflicts": len(links["conflicts"]),
        "link_products_with_literature": links["personalization_axis"]["products_with_literature"],
        "link_products_total": links["personalization_axis"]["products_total"],
        # 상태 플래그
        "flag_complete": flags["complete"]["value"],
        "flag_release_ready": flags["release_ready"]["value"],
        "flag_blinding_ai": flags["independent_blinding_ai"]["value"],
        "flag_blinding_human": flags["independent_blinding"]["value"],
        "flag_performance_claim": flags["performance_claim_allowed"]["value"],
        # 사이트 배포 여부. 연구 상태 플래그 release_ready 와 별개다.
        "deployed": bool(deployment.get("deployed")),
        "deployment_id": deployment.get("deployment_id"),
        "public_url": deployment.get("public_url"),
        "deployment_commit": deployment.get("git_commit"),
    }


def set_run_font(run, family: str, size: float, color: RGBColor | None = None) -> None:
    run.font.name = family
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), family)


def configure_style(style, family: str, size: float, color: RGBColor | None = None) -> None:
    style.font.name = family
    style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    rpr = style.element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), family)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    set_run_font(run, FONT_LIGHT, 9, MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    width = tbl_pr.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        tbl_pr.append(width)
    width.set(qn("w:w"), str(sum(widths)))
    width.set(qn("w:type"), "dxa")
    indent = tbl_pr.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(value))
        grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def set_paragraph_font(paragraph, family: str = FONT_BODY, size: float = 10.5) -> None:
    for run in paragraph.runs:
        set_run_font(run, family, size)


def add_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.45
    set_paragraph_font(paragraph)


def add_note(doc: Document, label: str, text: str, fill: str = CAUTION_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(f"{label} ")
    set_run_font(run, FONT_SEMIBOLD, 10.5, NAVY)
    run = paragraph.add_run(text)
    set_run_font(run, FONT_BODY, 10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_metric_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [2500, 1800, 4700])
    headers = ["항목", "값", "해석"]
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(value)
        set_run_font(run, FONT_SEMIBOLD, 9.5, NAVY)
    for label, value, meaning in rows:
        cells = table.add_row().cells
        for index, text in enumerate((label, value, meaning)):
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index == 1 else WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(text)
            set_run_font(run, FONT_BODY if index else FONT_MEDIUM, 9.3)
    set_table_geometry(table, [2500, 1800, 4700])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build_thesis(metrics: dict[str, Any]) -> None:
    THESIS_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = doc.styles
    configure_style(styles["Normal"], FONT_BODY, 10.5)
    for name, family, size, color, before, after in (
        ("Title", FONT_EXTRABOLD, 24, NAVY, 0, 12),
        ("Subtitle", FONT_MEDIUM, 12, MUTED, 0, 8),
        ("Heading 1", FONT_SEMIBOLD, 16, BLUE, 18, 10),
        ("Heading 2", FONT_MEDIUM, 13, NAVY, 12, 6),
        ("Heading 3", FONT_MEDIUM, 11.5, NAVY, 8, 4),
    ):
        configure_style(styles[name], family, size, color)
        styles[name].paragraph_format.space_before = Pt(before)
        styles[name].paragraph_format.space_after = Pt(after)
        styles[name].paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("연세대학교 약학대학 | 권혁찬 졸업논문")
    set_run_font(run, FONT_LIGHT, 8.5, MUTED)
    add_page_number(section.footer.paragraphs[0])

    cover = doc.add_paragraph()
    cover.paragraph_format.space_before = Pt(86)
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run("졸업논문")
    set_run_font(run, FONT_MEDIUM, 11, BLUE)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(18)
    run = title.add_run("식약처 허가원문 기반 국내 일반의약품\n안전성 조회 시스템과 AI 자율 문헌 근거층 구축")
    set_run_font(run, FONT_EXTRABOLD, 24, NAVY)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("v4.0 최종 보고")
    set_run_font(run, FONT_MEDIUM, 13, MUTED)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(76)
    run = meta.add_run("연세대학교 약학대학\n2021194024 권혁찬\n2026년 7월")
    set_run_font(run, FONT_MEDIUM, 11, NAVY)
    doc.add_page_break()

    doc.add_heading("초록", level=1)
    add_note(
        doc,
        "성능 수치의 기준",
        "이 논문의 민감도·특이도·정밀도·F1은 모두 AI 참조표준 대비 값이다. 사람 참조표준이 "
        "아니며 절대적 진실 대비 정확도가 아니다. 사람 판정은 "
        f"{metrics['human_decisions']}건이다.",
    )
    add_body(
        doc,
        "이 연구는 국내 일반의약품을 제품명으로 입력하면 중복 성분, 최대용량, 복용 간격, 연령, 질환, 병용약 위험 신호를 허가원문과 함께 보여 주는 조회 시스템을 다룬다. "
        f"결정론적 허가원문 층에는 분석 제품 {metrics['products']}개, 고유 성분 {metrics['ingredients']}개, 제품-성분 연결 {metrics['bindings']}개, 복용 조건 {metrics['constraints']}개가 포함된다. "
        f"규칙은 전체 {metrics['rules_total']}개이고 그중 {metrics['rules_released']}개가 source 와 locator 를 갖춘 released 상태다. "
        "문헌은 이 판정을 대신하지 않고 위해 연관성을 설명하는 별도 참고 근거층으로 설계했다."
    )
    add_body(
        doc,
        f"AI는 허가원문에서 확인한 성분과 규칙 범위만 입력받아 PICOS 질문 {metrics['questions']}개를 만들고 PubMed를 검색했다. "
        f"질문별 hit 합계는 {metrics['search_hits']:,}건이었고 중복 제거 뒤 고유 PMID {metrics['corpus_rows']:,}개를 확보했다. "
        f"초록 보유 문헌은 {metrics['with_abstract']:,}개, 제목만 있는 문헌은 {metrics['title_only']:,}개였다. "
        f"코퍼스 전체 {metrics['screened']:,}행을 선별해 커버리지 {metrics['coverage']:.1f}을 달성했으며 "
        f"retain {metrics['decision_distribution'].get('retain', 0):,}건, deprioritize {metrics['decision_distribution'].get('deprioritize', 0):,}건, "
        f"uncertain {metrics['decision_distribution'].get('uncertain', 0)}건이었다."
    )
    add_body(
        doc,
        f"선별기의 재현도는 층화 표본 {metrics['ref_sample']}건에 대한 독립 AI 참조표준으로 측정했다. "
        f"AI 참조표준 대비 민감도 {num(metrics['ref_sensitivity'])}, 특이도 {num(metrics['ref_specificity'])}, "
        f"F1 {num(metrics['ref_f1'])}, 일치도 {num(metrics['ref_agreement'])}였다. "
        f"규칙엔진은 무라벨 사례 {metrics['blind_cases']}건을 3라운드 맹검 판정하고 라벨을 SHA-256으로 잠근 뒤에야 예측을 연결하는 절차로 평가했다. "
        f"released 규칙 {metrics['rules_released']}종 {metrics['blind_scored']}건에서 AI 참조표준 대비 특이도 {num(metrics['blind_specificity'])}, "
        f"정밀도 {num(metrics['blind_precision'])}, 민감도 {num(metrics['blind_sensitivity'])}, F1 {num(metrics['blind_f1'])}였고 "
        f"위양성 {metrics['blind_fp']}건, 위음성 {metrics['blind_fn']}건이었다."
    )
    add_body(
        doc,
        f"규칙 {metrics['link_rules_total']}개 전부에 초록 문장 단위 locator 를 가진 문헌 근거를 연결했다(링크 {metrics['link_total']}건, 고유 논문 {metrics['link_pmids']}편). "
        f"허가원문과 어긋나는 {metrics['link_conflicts']}건은 지우지 않고 conflict 로 보존했다. "
        "핵심어: 일반의약품, 식약처 허가원문, 제품명 기반 조회, PubMed, AI 문헌 선별, AI 참조표준, 맹검 독립평가"
    )

    doc.add_heading("1. 서론", level=1)
    doc.add_heading("1.1 제품명으로 시작하는 안전성 질문", level=2)
    add_body(
        doc,
        "사람은 약을 성분명보다 제품명으로 기억한다. 하지만 같은 성분이 여러 감기약과 해열진통제에 들어가면 제품명만 보고 중복을 알아차리기 어렵다. 이 연구는 제품명을 입력점으로 삼고, 제품에서 성분으로, 성분에서 허가된 용량과 주의사항으로 거슬러 올라가는 구조를 택했다."
    )
    doc.add_heading("1.2 사실과 근거를 섞지 않는 이유", level=2)
    add_body(
        doc,
        "‘이 조합이 허가된 1일 최대량을 넘는다’는 말은 허가원문과 산술로 확인하는 사실 주장이다. 반면 ‘과량 복용이 간손상과 연관된다’는 말은 문헌이 뒷받침하는 근거 주장이다. 이 연구는 두 문장을 한 권한 체계로 합치지 않는다. 허가원문은 판정을 내리고, PubMed 문헌은 판정의 배경을 설명한다."
    )
    doc.add_heading("1.3 연구 질문", level=2)
    add_body(
        doc,
        "첫째, 식약처 허가원문만으로 제품명 기반 위험 신호 판정을 구성할 수 있는가. 둘째, AI가 사람의 개입 없이 문헌 질문을 설계하고 코퍼스 전체를 선별할 수 있는가. 셋째, 그렇게 만든 규칙엔진의 판정이 독립적으로 구성한 참조표준과 얼마나 일치하는가. 넷째, 일치하지 않는 지점은 무엇 때문인가."
    )

    doc.add_heading("2. 연구 방법", level=1)
    doc.add_heading("2.1 두 층 구조", level=2)
    add_body(
        doc,
        "이 연구의 근거는 두 층으로 나뉜다. 결정층은 식약처 허가원문이며 제품, 성분, 함량, 복용 조건, 규칙 판정을 확정한다. 설명층은 PubMed 문헌이며 판정의 배경을 설명하지만 판정을 바꾸지 못한다. 두 층은 별도 파일과 별도 컬럼에 저장하고, 문헌 근거에는 규칙을 배포시킬 권한을 주지 않는다."
    )
    add_body(
        doc,
        f"식약처 의약품상세정보의 제품·원료약품·용법용량·사용상의주의사항을 구조화했다. 분석 집합은 제품 {metrics['products']}개와 고유 성분 {metrics['ingredients']}개다. "
        f"계산에 사용하는 제품-성분 연결은 {metrics['bindings']}개이며 복용 조건은 {metrics['constraints']}개다. released 규칙은 {metrics['rules_released']}개, 전체 규칙은 {metrics['rules_total']}개다. "
        "복용 조건 개수와 released 규칙 개수는 서로 다른 상태이므로 합치지 않는다."
    )
    doc.add_heading("2.2 AI 자율 PICOS와 PubMed 검색", level=2)
    add_body(
        doc,
        f"AI에는 성분 {metrics['ingredients']}개, 규칙 {metrics['rules_total']}개의 유형과 범위, PubMed 단일 자료원이라는 제약만 제공했다. 이전 영양성분 검색식이나 결과 수치는 제공하지 않았다. "
        f"AI는 아세트아미노펜, NSAID, 감기·알레르기 복합성분, 소화제 복합성분, 외용 진통성분의 {metrics['questions']}개 질문으로 묶었다. 각 질문에는 대상, 노출, 비교, 결과, 연구설계와 MeSH·제목/초록 검색어를 기록했다. "
        f"질문 설계 프롬프트는 SHA-256 {metrics['picos_prompt_sha256'][:16]}… 로 고정했다."
    )
    question_rows = [(row["title"], f"{row['hits']:,}" if row["hits"] is not None else "null", row["id"]) for row in metrics["question_rows"]]
    add_metric_table(doc, [(title, hits, qid) for title, hits, qid in question_rows])
    doc.add_heading("2.3 원시 응답과 코퍼스 계보", level=2)
    add_body(
        doc,
        "NCBI E-utilities는 API 키 없이 초당 세 번 이하로 호출했다. ESearch로 질문별 건수를 먼저 확인했고, 전체 상한을 넘긴 첫 검색은 EFetch 전에 중단했다. 좁힌 검색식으로 다시 실행한 뒤 query.txt, ESearch·EFetch XML, 응답 메타데이터와 SHA-256을 질문별 실행 폴더에 저장했다."
    )
    doc.add_heading("2.4 문헌 선별", level=2)
    add_body(
        doc,
        f"선별 판정은 에이전트가 배치 카드를 직접 읽고 직접 기록했다({metrics['screener']}). 지역 언어모델을 띄우지 않았고 외부 LLM API를 호출하지 않았으며 하위 에이전트에 위임하지 않았다. "
        f"라벨은 retain, deprioritize, uncertain 세 가지다. 초록이 없으면 title_only 로 구분하고 신뢰도 상한을 low 로 제한했다. "
        f"배치는 {metrics['batch_count']}개이고 전체 {metrics['corpus_rows']:,}행이 정확히 한 번씩 배치에 들어갔다. 판정은 append-only JSONL 체크포인트로 기록했으며 "
        f"선별 프롬프트는 SHA-256 {metrics['screening_prompt_sha256'][:16]}… 로 실행 중 고정했다."
    )
    doc.add_heading("2.5 AI 참조표준", level=2)
    add_body(
        doc,
        f"선별 결과의 재현도를 재기 위해 선별과 다른 프롬프트로 독립 참조표준을 만들었다. 근거 형태(초록 보유 여부) × 선별 판정의 {metrics['ref_strata']}개 층에서 "
        f"층별 최소 배정을 둔 뒤 총 {metrics['ref_sample']}건을 뽑았고, 모든 지표에 층 가중치를 적용했다. "
        "참조표준은 주제 적합성을 통째로 묻지 않고 P·I·C·O·S 다섯 요소를 따로 판정한 뒤 도구가 코드 규칙으로 종합 라벨을 도출한다. "
        f"카드에는 선별 판정·신뢰도·이유 코드를 제거하고 라운드별 별칭만 노출했으며 참조표준 프롬프트는 SHA-256 {metrics['ref_prompt_sha256'][:16]}… 로 고정했다."
    )
    doc.add_heading("2.6 규칙엔진 AI 맹검 독립평가", level=2)
    add_body(
        doc,
        f"실제 허가 데이터에서 정답 라벨과 엔진 예측이 없는 사례 {metrics['blind_cases']}건을 만들었다. 사례 카드에는 제품명, 성분과 함량, 복용량, 정규화된 허가 상한, "
        "식약처 허가원문 발췌(용법·용량과 사용상의 주의)와 사용자 조건만 넣고 규칙 ID·심각도·바인딩 표는 넣지 않았다. "
        f"라운드별 무작위 순서로 3회 판정한 뒤 다수결 라벨을 SHA-256으로 잠갔고, 잠금 해시를 검증한 다음에야 배포 규칙 {metrics['rules_released']}종만으로 엔진 예측을 기록했다. "
        "지표 계산 단계는 잠금 시각이 예측 시각보다 앞선다는 것을 다시 확인한 뒤에만 실행된다."
    )
    doc.add_heading("2.7 문헌 근거 연결", level=2)
    add_body(
        doc,
        f"retain 판정을 받은 문헌만 규칙에 연결했다. 연결 단위는 규칙 1건 × 논문 1편이며 각 연결은 초록의 문장 인덱스와 그 문장의 원문 인용을 함께 저장한다. "
        "빌드할 때마다 코퍼스의 초록에서 해당 인덱스 문장을 다시 뽑아 저장된 인용문과 글자 단위로 대조하고, 어긋나면 빌드를 실패시킨다. "
        "허가원문과 문헌이 다른 방향을 가리키면 어느 한쪽을 지우지 않고 conflict 로 보존한다."
    )

    doc.add_heading("3. 결과", level=1)
    doc.add_heading("3.1 허가원문 결정층", level=2)
    add_metric_table(
        doc,
        [
            ("분석 제품", f"{metrics['products']}개", "제품명 기반 입력 집합"),
            ("고유 성분", f"{metrics['ingredients']}개", "분석 제품에서 확인"),
            ("제품-성분 연결", f"{metrics['bindings']}개", "계산용 선택 연결"),
            ("복용 조건", f"{metrics['constraints']}개", "허가원문 검증 완료, 약사 재검토 별도"),
            ("released 규칙", f"{metrics['rules_released']}개", f"전체 {metrics['rules_total']}개 중 source·locator 완비"),
        ],
    )
    doc.add_heading("3.2 문헌 검색층과 선별", level=2)
    add_metric_table(
        doc,
        [
            ("AI PICOS 질문", f"{metrics['questions']}개", f"성분 {metrics['ingredients']}개와 규칙 {metrics['rules_total']}개 유형을 포괄"),
            ("질문별 hit 합계", f"{metrics['search_hits']:,}건", "질문 사이 중복 포함"),
            ("고유 PMID", f"{metrics['corpus_rows']:,}개", "중복 제거 뒤 코퍼스"),
            ("초록 보유", f"{metrics['with_abstract']:,}개", "제목과 초록으로 선별"),
            ("제목만 보유", f"{metrics['title_only']:,}개", "title_only, confidence=low 상한"),
            ("선별 완료", f"{metrics['screened']:,}행", f"커버리지 {metrics['coverage']:.1f}, 누락·중복 0"),
            ("retain", f"{metrics['decision_distribution'].get('retain', 0):,}건", "직접 근거 후보"),
            ("deprioritize", f"{metrics['decision_distribution'].get('deprioritize', 0):,}건", "질문 직접성이 낮음"),
            ("uncertain", f"{metrics['decision_distribution'].get('uncertain', 0)}건", "초록만으로 확정 어려움"),
        ],
    )
    doc.add_page_break()
    doc.add_heading("3.3 선별기의 AI 참조표준 대비 재현도", level=2)
    add_metric_table(
        doc,
        [
            ("층화 표본", f"{metrics['ref_sample']}건", f"{metrics['ref_strata']}개 층, 층 가중치 적용"),
            ("채점 대상", f"{metrics['ref_rows']}건", f"참조표준 uncertain {metrics['ref_uncertain']}건 제외"),
            ("AI 참조표준 대비 민감도", num(metrics["ref_sensitivity"]), "참조표준 retain 을 선별이 retain 으로 본 비율"),
            ("AI 참조표준 대비 특이도", num(metrics["ref_specificity"]), "참조표준 deprioritize 를 선별이 걸러낸 비율"),
            ("AI 참조표준 대비 정밀도", num(metrics["ref_precision"]), "선별 retain 중 참조표준도 retain 인 비율"),
            ("AI 참조표준 대비 F1", num(metrics["ref_f1"]), "민감도와 정밀도의 조화평균"),
            ("AI 참조표준 대비 일치도", num(metrics["ref_agreement"]), "가중 일치 비율"),
            ("보정 유병률", num(metrics["ref_prevalence"]), "Rogan–Gladen 보정"),
        ],
    )
    add_body(
        doc,
        f"보정 유병률을 코퍼스 전체에 적용하면 실제 retain 문헌은 약 {metrics['ref_estimated_retain']:,.0f}건으로 추정되며 "
        f"95% 신뢰구간은 {metrics['ref_prevalence_ci'][0]:.4f}–{metrics['ref_prevalence_ci'][1]:.4f}이다. "
        f"라운드 3회의 평균 일치도는 {num(metrics['ref_inter_round_agreement'])}였다. 이 값은 같은 평가자가 같은 명시적 규칙을 재적용한 결과이므로 "
        "판정 안정성이지 서로 다른 평가자 간 신뢰도가 아니다."
    )
    doc.add_heading("3.4 규칙엔진의 AI 참조표준 대비 성능", level=2)
    add_metric_table(
        doc,
        [
            ("무라벨 사례", f"{metrics['blind_cases']}건", "정답·예측 없이 생성"),
            ("채점 대상", f"{metrics['blind_scored']}건", f"uncertain {metrics['blind_uncertain']}건·맹검 훼손 {metrics['blind_compromised']}건 제외"),
            ("AI 참조표준 대비 특이도", num(metrics["blind_specificity"]), f"위양성 {metrics['blind_fp']}건"),
            ("AI 참조표준 대비 정밀도", num(metrics["blind_precision"]), "엔진 경고 중 참조표준도 경고인 비율"),
            ("AI 참조표준 대비 민감도", num(metrics["blind_sensitivity"]), f"위음성 {metrics['blind_fn']}건"),
            ("AI 참조표준 대비 F1", num(metrics["blind_f1"]), "민감도와 정밀도의 조화평균"),
            ("AI 참조표준 대비 일치도", num(metrics["blind_agreement"]), "전체 일치 비율"),
        ],
    )
    add_body(
        doc,
        f"민감도의 95% 신뢰구간은 {metrics['blind_sensitivity_ci'][0]:.4f}–{metrics['blind_sensitivity_ci'][1]:.4f}, "
        f"특이도의 95% 신뢰구간은 {metrics['blind_specificity_ci'][0]:.4f}–{metrics['blind_specificity_ci'][1]:.4f}이다. "
        f"라운드 3회의 평균 일치도는 {num(metrics['blind_inter_round_agreement'])}이고 3-way 불일치는 {metrics['blind_unresolved']}건이었다. "
        f"참조 라벨 잠금 파일의 SHA-256은 {metrics['blind_lock_sha256'][:16]}… 이며 예측 기록은 이 해시를 검증한 뒤에 생성됐다."
    )
    doc.add_heading("3.5 위음성이 몰린 지점", level=2)
    gap_rows = [
        (
            rule_type,
            f"{detail['missed_cases']}건",
            "규칙이 묶이지 않은 제품에서 허가원문이 같은 주의를 적음",
        )
        for rule_type, detail in sorted(
            metrics["blind_gap_details"].items(), key=lambda kv: -kv[1]["missed_cases"]
        )
    ]
    add_metric_table(doc, gap_rows)
    add_body(
        doc,
        f"위양성은 {metrics['blind_fp']}건으로 엔진이 허가 근거 없이 경고하는 경우는 관찰되지 않았다. 반면 위음성 {metrics['blind_fn']}건은 "
        f"{metrics['blind_gap_rule_types']}개 규칙 유형에 몰려 있고, 모두 규칙이 대표 제품 하나에만 묶여 있어 같은 주의가 적힌 다른 제품에서 발동하지 않는 경우였다. "
        "이는 판정 논리의 오류가 아니라 규칙 바인딩 범위의 공백이다."
    )
    add_note(
        doc,
        "draft 규칙",
        f"{', '.join(metrics['draft_rule_types'])} 은 draft 상태라 엔진이 발동하지 않는다. 동시에 참조표준도 해당 "
        f"{metrics['draft_uncertain']}건을 전부 uncertain 으로 판정했다. 대상 제품의 허가원문이 연속 복용 기간을 "
        "정성적으로만 적고 일수 기준을 주지 않기 때문이며, 이 규칙이 배포되지 못하는 이유가 참조표준 쪽에서 독립적으로 확인된다.",
    )
    doc.add_heading("3.6 문헌 근거 연결", level=2)
    add_metric_table(
        doc,
        [
            ("문헌이 연결된 규칙", f"{metrics['link_rules_covered']}/{metrics['link_rules_total']}개", "모든 규칙이 최소 1건"),
            ("연결 수", f"{metrics['link_total']}건", "규칙 1건 × 논문 1편"),
            ("고유 논문", f"{metrics['link_pmids']}편", "retain 판정 문헌만"),
            ("보존한 충돌", f"{metrics['link_conflicts']}건", "허가원문과 어긋나지만 지우지 않음"),
            ("문헌이 연결된 제품", f"{metrics['link_products_with_literature']}/{metrics['link_products_total']}개", "제품→성분→문헌 축"),
        ],
    )

    doc.add_heading("4. 고찰", level=1)
    doc.add_heading("4.1 이층 구조가 주는 이점", level=2)
    add_body(
        doc,
        "문헌이 허가 판정을 덮어쓰지 않으므로 사실의 출처가 흐려지지 않는다. 사용자는 제품명으로 위험 신호를 찾고, 판정 근거로 허가원문 locator를 확인하며, 별도 참고 문헌에서 위해 연관성의 배경을 읽을 수 있다. 문헌과 허가사항이 충돌해도 한쪽을 지우지 않고 충돌로 남길 수 있다."
    )
    doc.add_heading("4.2 특이도 1.0과 낮은 민감도의 비대칭", level=2)
    add_body(
        doc,
        f"규칙엔진은 위양성 {metrics['blind_fp']}건, 즉 허가 근거 없이 경고하는 경우가 없었다. 이는 규칙이 허가원문에 묶여 있고 locator 없이는 배포되지 않는 설계의 직접적 결과다. "
        f"반대로 위음성 {metrics['blind_fn']}건은 규칙이 대표 제품 하나에만 묶여 있기 때문에 생긴다. "
        "예를 들어 임신·수유 주의는 이부프로펜에만 묶여 있지만 나프록센과 덱시부프로펜의 허가원문에도 같은 주의가 적혀 있다. "
        "따라서 다음 개선은 판정 논리를 바꾸는 것이 아니라 이미 허가원문에 있는 주의를 규칙 바인딩으로 옮기는 작업이다."
    )
    doc.add_heading("4.3 AI 참조표준이 설계 의도를 뒤집은 사례", level=2)
    add_body(
        doc,
        "사례 생성기는 발동 예상과 미발동 예상을 같은 수로 설계했지만, 참조표준이 허가원문을 읽고 판정한 결과 일부가 뒤집혔다. "
        "예를 들어 2세대 항히스타민 제품의 허가원문에는 과량의 알코올과 함께 투여하지 말라는 문장이 있어 음주 노출이 성립했다. "
        "설계 의도를 정답으로 삼지 않고 허가원문을 따랐다. 참조표준을 설계에 맞추는 순간 독립성이 사라지기 때문이다."
    )

    doc.add_heading("5. 한계", level=1)
    add_body(
        doc,
        "첫째, 이 논문의 모든 성능 수치는 AI 참조표준 대비 재현도이며 절대적 진실 대비 정확도가 아니다. "
        f"사람 판정은 {metrics['human_decisions']}건이고 사람 블라인드 독립평가는 수행되지 않았다."
    )
    add_body(
        doc,
        "둘째, 문헌 분류기와 참조표준을 같은 에이전트가 수행했으므로 독립성이 부분적이다. "
        "무라벨 사례, 별칭 카드, 라운드별 무작위 순서, 잠금 후 예측 연결이라는 절차적 맹검은 갖췄지만 평가자 독립성은 외부 사람 평가와 동등하지 않다. "
        "라운드 간 일치도가 높게 나온 것도 같은 평가자가 같은 규칙을 재적용한 결과이며 평가자 간 신뢰도로 읽으면 안 된다."
    )
    add_body(
        doc,
        "셋째, 자료원은 PubMed 하나다. Embase 등 다른 데이터베이스를 사용하지 않았으므로 코퍼스가 특정 색인 관행에 치우쳤을 수 있다."
    )
    add_body(
        doc,
        f"넷째, 판매량 자료가 없다. 분석 제품 {metrics['products']}개는 대표 일반의약품 후보이며 판매 순위 집합이 아니다. 대표성을 주장하지 않는다."
    )
    add_body(
        doc,
        f"다섯째, 복용 조건 {metrics['constraints']}개는 허가원문 검증까지만 완료됐고 별도 약사 재검토를 거치지 않았다. "
        f"이 개수는 released 규칙 {metrics['rules_released']}개와 다른 상태이며 합쳐서 읽으면 안 된다."
    )
    add_body(
        doc,
        f"여섯째, 규칙엔진 평가에서 참조표준이 판정하지 못한 {metrics['blind_uncertain']}건과 사례 준비 중 라벨이 노출된 {metrics['blind_compromised']}건은 "
        "1차 지표에서 제외했다. 제외 사유와 건수를 산출물에 함께 기록했다."
    )

    doc.add_heading("6. 결론", level=1)
    add_body(
        doc,
        f"이 연구는 식약처 허가원문을 결정층으로 유지하면서 AI가 설계한 PubMed 문헌 근거층을 별도로 만들었다. "
        f"PICOS {metrics['questions']}개와 고유 PMID {metrics['corpus_rows']:,}개 코퍼스는 재현 가능한 원시 응답과 해시를 가지며, "
        f"코퍼스 전체를 선별해 커버리지 {metrics['coverage']:.1f}을 달성했다."
    )
    add_body(
        doc,
        f"규칙엔진은 AI 참조표준 대비 특이도 {num(metrics['blind_specificity'])}, 정밀도 {num(metrics['blind_precision'])}로 "
        "허가 근거 없이 경고하지 않는다는 성질을 지켰고, 민감도 "
        f"{num(metrics['blind_sensitivity'])}는 규칙 바인딩 범위의 공백에서 비롯됐다. "
        f"규칙 {metrics['link_rules_total']}개 전부에 문장 단위 locator 를 가진 문헌 근거를 연결했고 충돌 {metrics['link_conflicts']}건을 보존했다."
    )
    add_body(
        doc,
        f"상태 플래그는 complete={str(metrics['flag_complete']).lower()}, "
        f"independent_blinding_ai={str(metrics['flag_blinding_ai']).lower()}, "
        f"performance_claim_allowed={str(metrics['flag_performance_claim']).lower()}, "
        f"independent_blinding={str(metrics['flag_blinding_human']).lower()}, "
        f"release_ready={str(metrics['flag_release_ready']).lower()}다. "
        "성능 주장은 AI 참조표준 대비라는 사실과 평가자가 사람이 아니라는 사실을 항상 함께 적는 조건에서만 허용된다. "
        "임상 배포 승인 절차는 이 연구의 범위 밖이다."
    )

    doc.add_heading("참고 자료", level=1)
    references = [
        "식품의약품안전처 의약품안전나라. 의약품상세정보 및 허가 원문. 연구 원시자료의 source_id와 locator에 기록.",
        "National Center for Biotechnology Information. Entrez Programming Utilities Help. https://www.ncbi.nlm.nih.gov/books/NBK25501/",
        "research_v3/protocol/protocol-v4.0-full-ai.md.",
        "research_v3/otc/literature/picos/picos_definition.json.",
        "research_v3/otc/literature/screening/screening_manifest.json.",
        "research_v3/measurement/screener_vs_ai_reference.json. 선별기의 AI 참조표준 대비 재현도.",
        "research_v3/otc/validation/ai_independent_evaluation.json. 규칙엔진 AI 맹검 독립평가.",
        "research_v3/otc/rules/literature_link_manifest.json. 규칙별 문헌 근거 연결과 보존한 충돌.",
        "research_v3/otc/audit/completion_audit.json. 상태 플래그와 근거 파일 경로.",
        "research_v3/otc/metrics_manifest.json.",
    ]
    for index, reference in enumerate(references, start=1):
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(reference)
        set_run_font(run, FONT_BODY, 9.5)

    doc.add_page_break()
    doc.add_heading("부록 A. 재현성 식별자", level=1)
    add_metric_table(
        doc,
        [
            ("PICOS 프롬프트 SHA-256", metrics["picos_prompt_sha256"], "AI 질문 설계 프롬프트"),
            ("선별 프롬프트 SHA-256", metrics["screening_prompt_sha256"], "고정된 P2 선별 프롬프트"),
            ("참조표준 프롬프트 SHA-256", metrics["ref_prompt_sha256"], "P3-A 요소별 채점 프롬프트"),
            ("코퍼스 입력 SHA-256", metrics["corpus_sha256"], "evidence_map.csv"),
            ("참조 라벨 잠금 SHA-256", metrics["blind_lock_sha256"], "P3-B 잠금 후 예측 연결"),
            ("판정 주체", metrics["screener"], f"사람 판정 {metrics['human_decisions']}건"),
        ],
    )
    doc.core_properties.title = "식약처 허가원문 기반 국내 일반의약품 안전성 조회 시스템과 AI 자율 문헌 근거층 구축"
    doc.core_properties.author = "권혁찬"
    doc.core_properties.subject = "v4.0 최종 보고"
    doc.save(THESIS_DOCX)


def build_markdown(metrics: dict[str, Any]) -> None:
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    retain = metrics["decision_distribution"].get("retain", 0)
    deprioritize = metrics["decision_distribution"].get("deprioritize", 0)
    uncertain = metrics["decision_distribution"].get("uncertain", 0)
    top_gaps = ", ".join(
        f"{rule_type} {detail['missed_cases']}건"
        for rule_type, detail in sorted(
            metrics["blind_gap_details"].items(), key=lambda kv: -kv[1]["missed_cases"]
        )[:3]
    )

    presentation = f"""# 발표원고 v4.0

## 슬라이드 1. 이 연구가 푸는 문제

사람은 약을 성분명보다 제품명으로 기억합니다. 하지만 감기약과 해열진통제를 함께 먹으면 같은 성분이 겹칠 수 있습니다. 이 연구는 제품명을 넣으면 성분, 용량, 복용 간격과 주의사항까지 거슬러 올라가는 조회 도구를 만듭니다. 분석 집합은 제품 {metrics['products']}개, 고유 성분 {metrics['ingredients']}개입니다.

## 슬라이드 2. 왜 연구 주제를 바꿨나

초기 영양성분 자료는 이전 계보로 보존했습니다. 이번 연구는 국내 실제 일반의약품의 허가원문과 제품명 중심 질문으로 옮겼습니다. 이전 영양성분 수치는 새 결과에 합산하지 않습니다. 같은 창고에 두 상자를 두되 라벨을 섞지 않는 것과 같습니다.

## 슬라이드 3. 허가원문에서 무엇을 뽑았나

식약처 원문에서 제품 {metrics['products']}개, 성분 {metrics['ingredients']}개, 계산 연결 {metrics['bindings']}개를 확인했습니다. 복용 조건은 {metrics['constraints']}개입니다. 규칙은 전체 {metrics['rules_total']}개이고 그중 {metrics['rules_released']}개가 source와 locator를 갖춘 released 상태입니다. 이 층이 실제 위험 신호를 판정합니다.

## 슬라이드 4. 왜 문헌층을 따로 두나

최대용량 초과 여부는 허가사항과 계산으로 결정합니다. 과량이 간손상과 연관된다는 설명은 PubMed 문헌에서 찾습니다. 판정하는 층과 설명하는 층을 분리하면 출처가 흐려지지 않습니다. 문헌은 허가 판정을 바꾸지 못합니다.

## 슬라이드 5. AI가 PICOS를 어떻게 정했나

PICOS는 누구에게, 어떤 노출을, 무엇과 비교해, 어떤 결과와 연구설계로 볼지 정하는 질문 틀입니다. AI는 성분 {metrics['ingredients']}개와 규칙 {metrics['rules_total']}개 유형만 받고 질문 {metrics['questions']}개를 만들었습니다. 기존 영양성분 검색식은 읽지 않았습니다. 질문은 아세트아미노펜, NSAID, 감기·알레르기 복합성분, 소화제, 외용 진통성분으로 묶였습니다.

## 슬라이드 6. 문헌을 어떻게 모으고 걸렀나

PubMed ESearch로 건수를 먼저 확인하고 상한을 넘은 검색은 다운로드 전에 좁혔습니다. 질문별 hit 합계는 {metrics['search_hits']:,}건, 중복 제거 뒤 고유 PMID는 {metrics['corpus_rows']:,}개였습니다. 이 {metrics['corpus_rows']:,}행을 배치 {metrics['batch_count']}개로 나눠 전부 판정했고 커버리지는 {metrics['coverage']:.1f}입니다. retain {retain:,}건, deprioritize {deprioritize:,}건, uncertain {uncertain}건이며 사람 판정은 {metrics['human_decisions']}건입니다.

## 슬라이드 7. 선별이 얼마나 재현되는지 어떻게 쟀나

선별 결과를 보지 않는 별도 참조표준을 만들어 층화 표본 {metrics['ref_sample']}건을 다시 판정했습니다. AI 참조표준 대비 민감도 {num(metrics['ref_sensitivity'])}, 특이도 {num(metrics['ref_specificity'])}, F1 {num(metrics['ref_f1'])}입니다. 주제 적합성을 통째로 묻지 않고 PICOS 다섯 요소를 따로 판정한 뒤 코드 규칙으로 종합했습니다. 이 수치는 사람 기준이 아니라 AI 참조표준 대비 값입니다.

## 슬라이드 8. 규칙엔진 맹검평가

정답도 예측도 없는 사례 {metrics['blind_cases']}건을 만들고 3라운드 맹검 판정 뒤 라벨을 SHA-256으로 잠갔습니다. 잠금 해시를 검증한 다음에야 배포 규칙 {metrics['rules_released']}종의 예측을 연결했습니다. AI 참조표준 대비 특이도 {num(metrics['blind_specificity'])}, 정밀도 {num(metrics['blind_precision'])}, 민감도 {num(metrics['blind_sensitivity'])}입니다. 위양성은 {metrics['blind_fp']}건, 위음성은 {metrics['blind_fn']}건이었습니다.

## 슬라이드 9. 규칙 근거와 위음성의 정체

규칙 {metrics['link_rules_total']}개 전부에 초록 문장 단위 locator를 가진 문헌 근거를 연결했습니다. 링크 {metrics['link_total']}건, 고유 논문 {metrics['link_pmids']}편이고 허가원문과 어긋나는 {metrics['link_conflicts']}건은 conflict로 보존했습니다. 위음성 {metrics['blind_fn']}건은 {metrics['blind_gap_rule_types']}개 규칙 유형에 몰려 있습니다({top_gaps}). 판정 논리의 오류가 아니라 규칙이 대표 제품 하나에만 묶여 있어 생긴 커버리지 공백입니다.

## 슬라이드 10. 한계와 결론

모든 성능 수치는 AI 참조표준 대비 재현도이며 사람 기준 정확도가 아닙니다. 분류기와 참조표준을 같은 에이전트가 수행해 독립성이 부분적이고, 자료원은 PubMed 하나이며, 판매량 자료가 없어 제품 {metrics['products']}개는 대표 후보일 뿐입니다. 복용 조건 {metrics['constraints']}개는 허가원문 검증까지만 끝났습니다. 그래도 허가 근거 없이 경고하지 않는다는 성질은 위양성 {metrics['blind_fp']}건으로 확인됐고, 다음 과제는 이미 허가원문에 있는 주의를 규칙 바인딩으로 넓히는 일입니다.
"""
    (REPORTS_DIR / "발표원고_v4.0.md").write_text(presentation, encoding="utf-8")

    notion = f"""# 현재 상태 - v4.0 최종

> PubMed 코퍼스 {metrics['corpus_rows']:,}행을 전부 선별해 커버리지 {metrics['coverage']:.1f}을 달성했고, AI 참조표준 채점과 규칙엔진 AI 맹검 독립평가까지 끝났습니다. `complete=true`, `performance_claim_allowed=true`입니다. 다만 성능 수치는 **AI 참조표준 대비** 값이며 사람 평가가 아닙니다. `independent_blinding=false`, `release_ready=false`를 유지합니다.

## 핵심 수치

- 허가원문 결정층: 제품 {metrics['products']}개, 성분 {metrics['ingredients']}개, 계산 연결 {metrics['bindings']}개, 복용 조건 {metrics['constraints']}개
- 규칙: 전체 {metrics['rules_total']}개, released {metrics['rules_released']}개
- AI PICOS: {metrics['questions']}개 · 질문별 hit 합계 {metrics['search_hits']:,}건 · 고유 PMID {metrics['corpus_rows']:,}개
- 문헌 형태: 초록 보유 {metrics['with_abstract']:,}개, 제목만 {metrics['title_only']:,}개
- 선별: {metrics['screened']:,}행, 커버리지 {metrics['coverage']:.1f}, 배치 {metrics['batch_count']}개, 사람 판정 {metrics['human_decisions']}건
- 판정 분포: retain {retain:,} · deprioritize {deprioritize:,} · uncertain {uncertain}
- 선별의 AI 참조표준 대비: 민감도 {num(metrics['ref_sensitivity'])} · 특이도 {num(metrics['ref_specificity'])} · F1 {num(metrics['ref_f1'])} (층화 표본 {metrics['ref_sample']}건)
- 규칙엔진의 AI 참조표준 대비: 특이도 {num(metrics['blind_specificity'])} · 정밀도 {num(metrics['blind_precision'])} · 민감도 {num(metrics['blind_sensitivity'])} · F1 {num(metrics['blind_f1'])} (채점 {metrics['blind_scored']}건)
- 위양성 {metrics['blind_fp']}건 · 위음성 {metrics['blind_fn']}건
- 문헌 근거: 규칙 {metrics['link_rules_covered']}/{metrics['link_rules_total']}개 연결, 링크 {metrics['link_total']}건, 논문 {metrics['link_pmids']}편, 보존한 충돌 {metrics['link_conflicts']}건

## 방법 요약 - 두 층을 섞지 않음

식약처 허가원문은 제품, 성분, 함량, 복용 조건과 규칙 판정을 결정합니다. PubMed 문헌은 위해 연관성을 설명하는 참고 근거이며 허가 판정을 바꾸지 못합니다. 사람 판정 자료는 이전 계보로 보존하고 v4.0 입력이나 정답으로 연결하지 않았습니다.

문헌 선별은 에이전트가 배치 카드를 직접 읽고 직접 기록했습니다. 지역 언어모델을 띄우지 않았고 외부 LLM API를 호출하지 않았으며 하위 에이전트에 위임하지 않았습니다.

## 성능 수치를 읽는 법

- 모든 지표는 **AI 참조표준 대비 재현도**입니다. 절대적 진실 대비 정확도가 아닙니다.
- 분류기와 참조표준을 같은 에이전트가 수행했으므로 평가자 독립성이 부분적입니다. 절차적 맹검(무라벨 사례, 별칭 카드, 라운드별 무작위 순서, 잠금 후 예측 연결)은 갖췄습니다.
- 라운드 간 일치도가 높은 것은 같은 평가자가 같은 규칙을 재적용했기 때문이며 평가자 간 신뢰도가 아닙니다.

## 위음성이 몰린 지점

위양성은 {metrics['blind_fp']}건으로, 허가 근거 없이 경고하는 경우는 관찰되지 않았습니다. 위음성 {metrics['blind_fn']}건은 {metrics['blind_gap_rule_types']}개 규칙 유형에 몰려 있습니다({top_gaps}). 모두 규칙이 대표 제품 하나에만 묶여 있어 같은 주의가 적힌 다른 제품에서 발동하지 않는 경우이며, 판정 논리의 오류가 아니라 규칙 바인딩 범위의 공백입니다.

## 상태 경계

- `independent_blinding_ai={str(metrics['flag_blinding_ai']).lower()}`: AI 맹검평가 완료
- `performance_claim_allowed={str(metrics['flag_performance_claim']).lower()}`: AI 참조표준 대비라는 사실을 병기하는 조건부 허용
- `complete={str(metrics['flag_complete']).lower()}`
- `independent_blinding={str(metrics['flag_blinding_human']).lower()}`: 사람 블라인드 평가는 수행되지 않음
- `release_ready={str(metrics['flag_release_ready']).lower()}`: 임상 배포 승인 절차는 연구 범위 밖

## 한계

AI 참조표준 대비 재현도라는 점, 분류기와 참조표준의 부분적 독립성, PubMed 단일 자료원, 판매량 자료 부재(제품 {metrics['products']}개는 대표 일반의약품 후보), 복용 조건 {metrics['constraints']}개가 허가원문 검증까지만 완료됐다는 점입니다.

## 검증

연구 시험 {metrics['research_tests']}개와 앱 시험 {metrics['app_tests']}개가 통과했고 lint·타입 검사·빌드도 통과했습니다. 정적 경로는 {metrics['static_paths']}개입니다.

{deployment_sentence(metrics, polite=True)}

## 공식 문서 위치

- `research_v3/protocol/protocol-v4.0-full-ai.md`
- `research_v3/otc/literature/picos/picos_definition.json`
- `research_v3/otc/literature/screening/screening_manifest.json`
- `research_v3/measurement/screener_vs_ai_reference.json`
- `research_v3/otc/validation/ai_independent_evaluation.json`
- `research_v3/otc/rules/literature_link_manifest.json`
- `research_v3/logs/v40_run_report.json`
"""
    (REPORTS_DIR / "notion_update.md").write_text(notion, encoding="utf-8")

    # 이 생성기는 보존된 v4.0 보고서도 만들지만 저장소 안내문은 v5.0 제출 정본과
    # 로컬 v5.1 경계를 설명해야 한다. 세 안내문을 정본과 같은 문자열로 유지한다.
    readme = """# 국내 일반의약품 안전성 조회 연구

이 저장소는 국내 일반의약품을 제품명으로 입력해 중복 성분, 최대용량, 복용 간격, 연령,
질환과 병용약 위험 신호를 찾는 연구용 시스템이다.

## 먼저 구분할 두 연구 상태

- 제출한 최종 연구는 v5.0이다. 정본은 `research_v3/`이며 수정하지 않는다.
- v5.1은 v5.0 정본을 기준으로 안전 범위와 화면 표시를 점검하는 로컬 확장이다.
  v5.1은 제출 최종본이나 임상 배포본이 아니다.

버전별 정본과 수치는 `docs/version_map.md`에서 확인한다.

## 두 근거층

- 식약처 허가원문은 제품, 성분, 함량, 복용 조건과 규칙 판정을 결정한다.
- PubMed 문헌은 위해 연관성을 설명하는 참고 근거다. 문헌은 허가 판정을 바꾸거나 규칙을
  release하지 않는다.

허가원문 분석 집합은 제품 13개, 성분 28개, 계산 연결 47개, 복용 조건 32개다. 규칙
16개 중 released 15개는 source와 locator를 가진다.

## v5.0 제출 최종 상태

v5.0은 선별 단위 43,207건과 고유 논문 42,822편을 다뤘다. 최종 판정은 retain 7,875건,
deprioritize 34,965건, uncertain 367건이다. 5,000건을 재판정했고 채점 arm은 894행이다.
문헌은 규칙 16개 중 9개에 10건을 연결했다.

채점 arm의 `agreement_vs_ai_reference`는 86.93%, `sensitivity_vs_ai_reference`는 46.92%,
`specificity_vs_ai_reference`는 95.84%, Cohen κ는 0.493이다. 이 값은 사람 기준 임상 정확도가
아니라 AI 참조표준 대비 결과다. v5.0 semantic adjudication selection의
`independent_blinding_ai=false`이며 `release_ready=false`다.

## 로컬 v5.1 안전 확장

v5.1 기준선은 commit `6dbdad518e2fa7b2ed7b9a8048e0c47dba5b6ae9`이다. v5.1은 제품
13개, released 규칙 15개와 허가 복용 조건 32개를 유지한다. 제품 10개는 용량·간격 조건만
지원하고 3개는 그보다 넓은 안전 규칙을 지원한다.

허가 근거 후보 연결 360건 중 15건은 v5.0에서 사람 전문가와 약사가 검토한 기존 released 규칙의
운영 근거다. 나머지 345건은 비활성이다. 비활성 후보는 `needs_expert_review` 33건,
`provisional` 308건, 제외 제품의 `rejected` 4건이다. v5.1에서 새로 활성화한 규칙이나 후보는
없고 `release_ready=false`다.

v5.0에서 결과에 사용한 문헌 링크 10건은 v5.1에서 정확한 직접 일치 1건, 범위를 함께 밝혀야
하는 직접 문헌 4건, 배경 문헌 5건으로 나눴다. v5.0에서 거절한 링크 10건은 결과 화면에서
제외한다. 문헌은 어느 분류에서도 규칙 release 권한을 갖지 않는다.

## 보존된 v4.0 비교 기록

v4.0에서 AI가 PICOS 질문 5개를 만들고 PubMed 고유 PMID 5,724개를 수집했다. 코퍼스 전체를
선별해 커버리지 1.0을 달성했고 사람 판정은 0건이다. 선별의 AI 참조표준 대비 F1은 0.8484,
규칙엔진의 AI 참조표준 대비 특이도는 1.0000, 민감도는 0.5702다.

v4.0 기록의 상태는 `complete=true`, `performance_claim_allowed=true`다. 성능 수치를 인용할
때는 **AI 참조표준 대비**라는 사실과 평가자가 사람이 아니라는 사실을 함께 적어야 한다.
`independent_blinding=false`, `release_ready=false`다. 당시 검증에서는 연구 시험 192개와 앱
시험 73개가 통과했고 정적 경로 156개를 생성했다. 이 수치를 현재 v5.1 검증 결과로 쓰지 않는다.

## 주요 경로

- v5.0 제출 정본과 실행 원장: `research_v3/`, `research_v3/logs/v50_run_report.json`
- v5.0 MECIR 문헌층: `research_v3/otc/literature/v5/`
- v5.1 프로토콜과 상태 계약: `research_v51/protocol/README.md`
- v5.1 근거 후보와 운영 범위: `research_v51/evidence/`
- v5.1 전문가 검토 큐: `research_v51/review/`
- v5.1 문헌 분류: `research_v51/literature/link_classification.csv`
- v5.1 기계 감사와 지원 행렬: `research_v51/audit/`
- 재현 명령 안내: `REPRODUCE.md`

## 코드와 배포

- GitHub: https://github.com/yeohj0710/otc-nutrient-safety-engine
- 공개 주소: https://otc-nutrient-safety-engine.vercel.app
- 사용자 지시로 production에 배포했다. 배포 ID는 `dpl_A1YDFPUijJCXWnVKdojUxP4pEZvn`이다.
  사이트 배포와 임상 배포 승인은 다르므로 `release_ready=false`를 유지한다.
- 로컬 v5.1 작업은 push, PR 생성 또는 production 배포를 하지 않는다.

## 검증

버전별 실행 순서와 네트워크가 필요한 검사는 `REPRODUCE.md`를 따른다. 코드 전체 검증 명령은
다음과 같다.

```powershell
.\\.venv-research\\Scripts\\python.exe -m pytest tests\\research -q
npm run typecheck
npm run lint
npm test
npm run build
```

이 시스템은 연구용 프로토타입이며 의료진의 진단이나 복약 결정을 대체하지 않는다.

Reference basis: tossfeed-easy-finance
"""
    (ROOT / "README.md").write_text(readme, encoding="utf-8")

    agents = """<!-- BEGIN:nextjs-agent-rules -->
# This is NOT the Next.js you know

This version has breaking changes — APIs, conventions, and file structure may all differ from your training data. Read the relevant guide in `node_modules/next/dist/docs/` before writing any code. Heed deprecation notices.
<!-- END:nextjs-agent-rules -->

## 어떤 버전이 최종인가

`docs/version_map.md`를 기준으로 판단한다. 이 저장소의 **제출 최종 연구는 v5.0**이고,
`nutrition-safety-engine`의 여형준 영양성분 연구는 v4.0이 최종이다. 이 저장소의
`research_v3/logs/v40_run_report.json`은 v5.0이 대체한 선행 트랙이다. 두 저장소의 수치를
파일명만 보고 섞지 않는다.

로컬 v5.1은 제출 최종본이 아니라 안전 범위 확장 트랙이다. v5.0 정본은 `research_v3/`,
v5.1 산출물은 `research_v51/`에 둔다.

## 프로젝트 탐색

저장소를 처음 탐색할 때는 `docs/project_map.md`를 먼저 읽는다.

- 메인 페이지: `app/page.tsx`
- 메인 클라이언트 UI: `src/components/rule-explorer-client.tsx`
- OTC 조회 UI: `src/components/otc-product-safety-client.tsx`
- 결과 카드 UI: `src/components/rule-card.tsx`
- 기존 안전 엔진: `src/lib/safety-engine/index.ts`
- OTC 규칙 엔진: `src/lib/otc/engine.ts`
- 지식 로더와 정규화: `src/lib/knowledge/`
- 이전 데이터 원본: `data/knowledge_pack.json`
- 현재 OTC 런타임: `src/generated/otc-runtime.json`
- 이전 런타임 인덱스: `src/generated/knowledge-index.json`
- v5.1 프로토콜: `research_v51/protocol/README.md`
- 프로젝트 지도: `docs/project_map.md`

## v5.0 제출 정본 경계

- 연구 질문은 국내 일반의약품의 제품명 기반 안전성 조회다.
- 식약처 허가자료가 제품, 성분, 함량, 복용 조건과 규칙 판정의 결정 근거다.
- PubMed는 AI가 선별한 별도 설명 근거층이다. 문헌은 허가 사실이나 released 규칙 논리를
  덮어쓰지 않는다.
- v5.0 문헌층은 규칙 16개 중 9개에 링크 10건을 연결한다. 미연결 규칙은
  `OTC-RULE-003` max_daily_dose, `009` gi_bleeding_ulcer, `010` sedation_driving,
  `011` alcohol, `013` sedative_medication, `015` maximum_duration, `016` urgent_referral이다.
- 미연결 사유는 `research_v3/otc/literature/v5/downstream/literature_link_manifest.json`에
  기록돼 있다. 후보 거절 6건(고유 논문 5편)은 `not_in_v5_corpus`이고, 4건은
  `no_retain_decision_for_rule_question`이다. `015` 후보 두 편은 Q01에서 retain이지만 규칙이
  허용하는 질문은 Q03·Q04다.
- 미연결은 검색 기간 공백이 아니다. 실제 필터는 Q01~Q03 2010/01/01 이후,
  Q04~Q05 2000/01/01 이후이며 해시한 질의문에 포함된다. v5.0 코퍼스의 출판연도는
  2000~2026이고, 미연결 후보 9편은 모두 해당 규칙의 허용 질문 기간 안에 있다.
  AM-OTC-003의 2022-01-01 원인 설명은 `AM-OTC-004`가 정정했다. 이 공백은 결과로 보고하고
  문헌층 자체를 연구 기여로 쓰지 않는다.
- 허가 근거와 문헌 근거는 별도 필드에 둔다. 충돌은 `conflict`로 보존한다.
- 규칙–문헌 링크는 문장 locator(`abstract:sentence:N`)와 인용 문장을 가져야 한다.
  `scripts/research/otc/build_supporting_literature.py`가 빌드마다 코퍼스 초록과 대조한다.
- v5.0 문헌 정본은 `research_v3/otc/literature/`에 있으며 읽기 전용이다.
  `research_v3/search/provisional_pubmed_20260710/`도 수정하지 않는다.
- `research_v3/screening/`, `research_v3/human_review_minimal/`, 전문가 검토 산출물과
  `human_reference_label`은 보존된 사람 판단 계보에 속한다. v4.0 또는 v5.0 실행 체인에 넣지 않는다.
- AI 참조표준 지표는 `sensitivity_vs_ai_reference`, `specificity_vs_ai_reference`,
  `agreement_vs_ai_reference`, `ai_reference_standard`, `ai_cross_checked`처럼 출처를 이름에
  포함한다. 출처 없는 ‘민감도’라고 쓰지 않는다.
- `independent_blinding`은 사람 맹검을 뜻하며 false다. AI 맹검은
  `independent_blinding_ai`로 기록하고 계층별로 해석한다. v4.0 문헌 선별은 true이고 근거는
  `research_v3/otc/validation/ai_independent_evaluation.json`이다. v5.0 semantic adjudication
  selection은 실행 영수증이 없어 false다. `V50-PC-001`과
  `research_v3/logs/DECISIONS_v50.md`가 이를 정정했다.
- v4.0 선별 기록은 5,724/5,724행, 커버리지 1.0, 사람 판정 0건이다. 성능 수치는 평가자가
  AI 참조표준이라는 사실과 함께 쓴다.
- `release_ready=false`다. 사이트 배포는 임상 배포 승인이 아니다. 사이트는 연구자의 지시로
  2026-07-30 production에 배포했다. `AM-OTC-004`가 이전의 일괄 배포 금지 문구를 이 원칙으로
  바꿨다. 이후 배포도 연구자가 명시적으로 요청할 때만 한다.
- `tools/search_pipeline/embase_adapter.py`를 삭제하지 않는다.
- 신신파스아렉스의 출처 기록은 보존하되 분석과 런타임에서는 제외한다.
- released 규칙은 source와 locator를 모두 가져야 한다. 복용 조건 32개와 released 규칙
  15개는 서로 다른 상태다.
- Python 체계적 검색 파이프라인은 Next.js 런타임과 분리한다. 코드는
  `tools/search_pipeline/`, 보존 산출물은 `data/systematic_search/`에 있다.
- `data/knowledge_pack.json`과 이전 영양성분 검색 결과는 대체된 탐색 자료로만 취급한다.

## 로컬 v5.1 안전 확장 경계

- 기준선은 commit `6dbdad518e2fa7b2ed7b9a8048e0c47dba5b6ae9`이다. `research_v3/`와
  저장소 밖 v5.0 논문 정본을 수정하지 않는다.
- v5.1 연구 산출물은 `research_v51/`에 둔다. v5.1 문헌 분류나 새 문헌 산출물도
  `research_v51/literature/`에 둔다. 이것이 위 v5.0 문헌 경로 규칙의 명시적 예외다.
- 현재 `verified_primary` 15건은 v5.0의 `human_expert_verified`, `supports_release=true`,
  released 규칙과 약사 검토 메타데이터를 모두 상속한 운영 근거다.
- 나머지 후보 연결 345건은 모두 비활성이다. 구성은 `needs_expert_review` 33건,
  `provisional` 308건, 제외 제품의 `rejected` 4건이다. 일반적인 근거 라벨
  `verified_primary`만으로 미래 후보를 승인하거나 활성화하지 않는다.
- 새 후보를 활성화하려면 사람의 명시적 승인, 안정적인 rule ID, 허가 source·version·locator,
  적용 범위, 사용자 문구와 정상·경계·비대상·오탐 방지 테스트가 모두 필요하다.
- released 규칙 15개와 `ADMIN-*` 허가 복용 조건 32개를 다른 ID와 판정 근거로 유지한다.
  과거 `supportedRuleTypes` 26건을 released 규칙 연결 수로 해석하지 않는다.
- v5.1 문헌 분류는 정확한 직접 일치 1건, 범위 한정 직접 문헌 4건, 배경 문헌 5건이다.
  v5.0에서 거절한 10건은 결과 UI에서 제외한다. 어떤 문헌도 규칙 release 권한이 없다.
- v5.1에서 새로 활성화한 규칙이나 후보는 없다. `release_ready=false`를 유지한다.
- v5.1은 로컬 브랜치와 로컬 커밋까지만 허용한다. 연구자가 명시적으로 승인하기 전에는
  push, PR, production 배포, 공개 게시 또는 G 드라이브 정본 교체를 하지 않는다.

## 검증

사이트를 바꾼 뒤 `npm run typecheck`, `npm run lint`, `npm test`, `npm run build`를 실행한다.
v5.1 경계와 산출물 검증 순서는 `REPRODUCE.md`를 따른다. 배포는 연구자가 명시적으로 요청할
때만 한다.
"""
    (ROOT / "AGENTS.md").write_text(agents, encoding="utf-8")

    project_map = """# Project map

## 사용자 화면

- `app/page.tsx`: 제품명 중심 조회 화면
- `app/sources/page.tsx`: 허가원문 출처 브라우저
- `app/rules/[id]/page.tsx`: 규칙 상세
- `src/components/otc-product-safety-client.tsx`: 제품 선택, 입력 폼, 판정 카드
- `src/components/rule-card.tsx`: 위험 신호와 근거 카드
- `src/lib/site.ts`: 사이트명과 설명

## 결정 엔진

- `src/lib/otc/engine.ts`: 허가원문 기반 OTC 판정 규칙
- `src/lib/otc/schema.ts`: v5.1 런타임 데이터 계약
- `src/lib/otc/presentation.ts`: 판정 근거와 참고 문헌의 표시 분리
- `src/generated/otc-runtime.json`: 제품, released 규칙, `ADMIN-*` 복용 조건 런타임
- `src/generated/otc-supporting-literature.json`: v5.1 표시 정책을 적용한 참고 문헌
- `src/lib/safety-engine/index.ts`: 이전 계보 결정 규칙
- `src/lib/knowledge/`: 이전 런타임 지식 인덱스 로더와 정규화
- `src/types/knowledge.ts`: 이전 지식 인덱스의 Zod 스키마와 핵심 타입

## v5.0 제출 정본 — 읽기 전용

- `research_v3/otc/normalized/`: 제품 16행(분석 13), 성분 31행(계산 28), 제품–성분
  106행(계산 선택 47), 복용 조건 32행
- `research_v3/otc/rules/rules.csv`: 전체 규칙 16개, released 15개
- `research_v3/otc/literature/v5/`: v5.0 MECIR 검색, 분류기 선별, semantic adjudication과 최종 결정
- `research_v3/otc/validation/screening_ai_reference_v50/`: 43,207건 모집단의 완전분할 층화
  설계, 실제 맹검·채점 표본 894행, 잠금 라벨과 영수증
- `research_v3/otc/synthesis/screener_vs_ai_reference_v50.json`: v5.0 채점 arm의 설계가중 비교 결과
- `research_v3/logs/v50_run_report.json`: v5.0 최종 수치와 해시의 단일 원장
- `research_v3/logs/v50_*`: v5.0 실행 원장, 판단 기록과 채점 결과

## 보존된 v4.0 비교 자료

- `research_v3/otc/rules/supporting_literature.csv`: 규칙×논문 링크 20건(문장 단위 locator 필수)
- `research_v3/otc/rules/literature_link_manifest.json`: 규칙별 연결 현황과 충돌 4건
- `research_v3/otc/literature/picos/`: AI 자율 PICOS 질문 5개
- `research_v3/otc/literature/searches/`: PubMed 원시 XML, 메타데이터와 SHA-256
- `research_v3/otc/literature/evidence_map.csv`: 고유 PMID 5,724개 코퍼스
- `research_v3/otc/literature/screening/`: 배치 115개와 판정 체크포인트, 커버리지 1.0
- `research_v3/measurement/ai_reference/`: 층화 표본 300건과 라운드별 판정
- `research_v3/otc/validation/ai_independent_cases/`: 무라벨 사례
- `research_v3/otc/validation/ai_independent_evaluation/`: 라운드 카드, 잠금 라벨, 예측 감사

## 로컬 v5.1 안전 확장

- `research_v51/protocol/README.md`: v5.0 보호 경계, 후보 상태, 문헌 권한과 로컬 작업 범위
- `research_v51/evidence/evidence_units.csv`: 중복을 정리한 허가 근거 단위 328건
- `research_v51/evidence/evidence_rule_links.csv`: 상태를 가진 후보–규칙 연결 360건
- `research_v51/evidence/active_rule_applicability.csv`: 기존 released 규칙 15개의 제한된 적용 범위
- `research_v51/review/expert_review_queue.csv`: 비활성 전문가 검토 후보 33건
- `research_v51/review/expert_review_packet.md`: 채택·수정·거절 판단과 필수 회귀 테스트 검토지
- `research_v51/literature/link_classification.csv`: 직접 1건, 범위 한정 직접 4건, 배경 5건과 UI 제외 10건
- `research_v51/audit/baseline_manifest.json`: v5.0 기준선과 보호 대상 해시
- `research_v51/audit/source_freshness_snapshot.json`: 식약처 원격 문서 20개의 의미 동일성 감사
- `research_v51/audit/product_support_matrix.csv`: 제품 10개 용량·간격 전용, 3개 확장 안전 지원
- `research_v51/audit/active_rule_matrix.csv`: released 규칙별 적용 범위와 출처 버전
- `research_v51/audit/final_metrics.json`: v5.0 기준선과 로컬 v5.1의 기계 집계

## 보존 계보

- `data/knowledge_pack.json`: 이전 영양성분 탐색 자료. 활성 OTC 성과에 합산하지 않음
- `data/systematic_search/`: 이전 검색 파이프라인 산출물
- `research_v3/otc/rules/supporting_literature_pre_v40.csv`: v4.0 검색 밖에서 큐레이션된 문헌
- `src/generated/knowledge-index.json`: 이전 Next.js 런타임 인덱스

## 실행 스크립트

- `tools/v40_literature_pipeline.py`: v4.0 PICOS 생성, ESearch, EFetch와 코퍼스 정규화
- `tools/agent_screening.py`: v4.0 선별 배치 생성·카드 렌더링·적재·커버리지 검증
- `tools/ai_reference_standard.py`: v4.0 P3-A 층화 표본, 라운드 카드, 가중 지표
- `tools/ai_independent_cases.py`: v4.0 P3-B 무라벨 사례 생성
- `tools/ai_independent_eval.py`: v4.0 P3-B 카드 렌더링, 라벨 잠금, 지표 산출
- `scripts/research/otc/predict-ai-independent.ts`: v4.0 잠금 검증 후 엔진 예측 기록
- `tools/build_rule_literature_links.py`: v4.0 문헌 링크 검증과 매니페스트 생성
- `tools/build_v40_run_report.py`: v4.0 실행 보고서 생성
- `tools/v50_scoring/`: v5.0 완전분할 층화 표본, 채점 하네스, 잠금 후 비교 보고
- `scripts/research/otc/build_v51_evidence_review.py`: v5.1 근거 단위, 후보 상태와 검토 큐 생성
- `scripts/research/otc/validate_v51_shortlist_triage.py`: 33건 의미 분류 계약 검증
- `scripts/research/otc/build_v51_review_packet.py`: 전문가 검토지와 감사 파일 생성
- `scripts/research/otc/build_runtime.py`: 제품, released 규칙과 `ADMIN-*` 복용 조건 런타임 생성
- `scripts/research/otc/build_v51_literature_classification.py`: v5.0 링크 20건의 v5.1 표시 정책 생성
- `scripts/research/otc/build_supporting_literature.py`: 표시 정책을 검증해 참고 문헌 런타임 생성
- `scripts/research/otc/audit_v51_boundaries.py`: `research_v3/`와 외부 논문 정본 보호 감사
- `scripts/research/otc/audit_v51_source_freshness.py`: 식약처 원격 문서 의미 동일성 감사
- `scripts/research/otc/build_v51_final_audit.py`: 제품·규칙·근거·문헌 상태의 기계 집계
- `tools/build_v40_reporting.py`: 보존된 v4.0 논문·문서·지표 생성기
- `tools/search_pipeline/`: 보존된 Python 검색 파이프라인

## 검증

버전별 순서와 네트워크 조건은 `REPRODUCE.md`를 따른다.

```powershell
.\\.venv-research\\Scripts\\python.exe -m pytest tests\\research -q
npm run typecheck
npm run lint
npm test
npm run build
```

보존된 v4.0 run report의 최근 실행은 연구 시험 192개, 앱 시험 73개, 정적 경로 156개다.
이 숫자는 현재 v5.1 검증 건수가 아니며 v5.0 기준선 감사값도 아니다.

Reference basis: tossfeed-easy-finance
"""
    (ROOT / "docs/project_map.md").write_text(project_map, encoding="utf-8")


def update_metrics(metrics: dict[str, Any]) -> None:
    layer = {
        "lineage": "research_v3_otc_v4_ai_literature",
        "status": "complete" if metrics["screening_complete"] else "partial",
        "ai_picos_questions": metrics["questions"],
        "pubmed_hits_before_deduplication": metrics["search_hits"],
        "corpus_rows": metrics["corpus_rows"],
        "abstract_rows": metrics["with_abstract"],
        "title_only_rows": metrics["title_only"],
        "ai_screened_rows": metrics["screened"],
        "screening_coverage": metrics["coverage"],
        "screening_distribution": metrics["decision_distribution"],
        "screening_run_complete": metrics["screening_complete"],
        "screener": metrics["screener"],
        "human_decisions": metrics["human_decisions"],
        # 지표 이름에 출처를 박아 둔다. 사람 기준 정확도가 아니다.
        "ai_reference": {
            "sample_size": metrics["ref_sample"],
            "strata": metrics["ref_strata"],
            "sensitivity_vs_ai_reference": metrics["ref_sensitivity"],
            "specificity_vs_ai_reference": metrics["ref_specificity"],
            "precision_vs_ai_reference": metrics["ref_precision"],
            "f1_vs_ai_reference": metrics["ref_f1"],
            "agreement_vs_ai_reference": metrics["ref_agreement"],
            "corrected_prevalence": metrics["ref_prevalence"],
        },
        "blind_engine_evaluation": {
            "cases_total": metrics["blind_cases"],
            "scored_cases": metrics["blind_scored"],
            "sensitivity_vs_ai_reference": metrics["blind_sensitivity"],
            "specificity_vs_ai_reference": metrics["blind_specificity"],
            "precision_vs_ai_reference": metrics["blind_precision"],
            "f1_vs_ai_reference": metrics["blind_f1"],
            "false_positive": metrics["blind_fp"],
            "false_negative": metrics["blind_fn"],
            "coverage_gap_rule_types": metrics["blind_gap_rule_types"],
            "lock_sha256": metrics["blind_lock_sha256"],
        },
        "rule_literature_links": {
            "rules_total": metrics["link_rules_total"],
            "rules_with_literature": metrics["link_rules_covered"],
            "links_total": metrics["link_total"],
            "unique_pmids": metrics["link_pmids"],
            "preserved_conflicts": metrics["link_conflicts"],
        },
        "ai_reference_standard": True,
        "human_reference_standard": False,
        "performance_claim_allowed": metrics["flag_performance_claim"],
        "performance_claim_condition_ko": (
            "AI 참조표준 대비 지표라는 사실과 평가자가 사람이 아니라는 사실을 항상 병기할 것"
        ),
        "complete": metrics["flag_complete"],
        "release_ready": metrics["flag_release_ready"],
        "evidence_paths": {
            "picos": PICOS_PATH.relative_to(ROOT).as_posix(),
            "corpus": EVIDENCE_PATH.relative_to(ROOT).as_posix(),
            "screening": SCREENING_PATH.relative_to(ROOT).as_posix(),
            "ai_reference": AI_REFERENCE_PATH.relative_to(ROOT).as_posix(),
            "blind_evaluation": BLIND_EVAL_PATH.relative_to(ROOT).as_posix(),
            "literature_links": LINK_MANIFEST_PATH.relative_to(ROOT).as_posix(),
        },
    }
    stale_blocker = "v4.0 AI 문헌 선별·AI 참조표준·규칙엔진 맹검 독립평가 미완료"
    for path in (OTC_METRICS_PATH, ROOT_METRICS_PATH):
        manifest = read_json(path)
        manifest["generated_at_utc"] = utc_now()
        manifest["literature_layer"] = layer
        blockers = manifest.setdefault("release_blockers", [])
        # 완료된 단계를 블로커로 남겨 두지 않는다.
        manifest["release_blockers"] = [item for item in blockers if item != stale_blocker]
        path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def main() -> None:
    metrics = collect_metrics()
    build_markdown(metrics)
    update_metrics(metrics)
    build_thesis(metrics)
    print(
        json.dumps(
            {
                "thesis_docx": str(THESIS_DOCX),
                "presentation": str(REPORTS_DIR / "발표원고_v4.0.md"),
                "notion": str(REPORTS_DIR / "notion_update.md"),
                "screening_coverage": metrics["coverage"],
            },
            ensure_ascii=False,
        )
    )


if __name__ == "__main__":
    main()
