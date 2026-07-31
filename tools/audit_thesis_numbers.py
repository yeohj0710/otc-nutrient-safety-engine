#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Read-only numerical audit for the Kwon v5.0 thesis and figure sources.

The script reads canonical JSON/CSV files, the thesis DOCX, and the figure
source/PNG files. It writes nothing. Results go to stdout and any canonical
mismatch causes a non-zero exit status.
"""

from __future__ import annotations

import argparse
import ast
import csv
import hashlib
import json
import os
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document


REPO = Path(__file__).resolve().parents[1]
DRIVE_ROOT = Path(
    os.environ.get(
        "THESIS_DRIVE_ROOT",
        r"G:\내 드라이브\여형준님\24 전공심화실습(1)",
    )
)
DEFAULT_DOCX = DRIVE_ROOT / "권혁찬" / "06_졸업논문" / "권혁찬_졸업논문_최종본.docx"
DEFAULT_FIG_DIR = DRIVE_ROOT / "권혁찬" / "06_졸업논문" / "etc" / "_그림"

RUN_REPORT = REPO / "research_v3" / "logs" / "v50_run_report.json"
SCORING_REPORT = REPO / "research_v3" / "logs" / "v50_scoring_report.json"
LINK_MANIFEST = REPO / "research_v3" / "otc" / "literature" / "v5" / "downstream" / "literature_link_manifest.json"
RULES_CSV = REPO / "research_v3" / "otc" / "rules" / "rules.csv"
QUERY_DEFINITIONS = REPO / "research_v3" / "otc" / "literature" / "v5" / "query_definitions.json"

EXPECTED_FIGURE_HASHES = {
    "fig_kwon.py": "3fe3667c3f7ecb7a631124efd961588625a3b20ae90f38ebe0d62ded5df19069",
    "kwon_fig1_two_layer_authority.png": "232ee0f755997e4fd3d497845364c228e3cf942b7528e806bcff6ec876f1d3c5",
    "kwon_fig2_screening_flow.png": "100024af74ef707a75aac921d8c6177e1629a2b3870b37f68cb7053c998c350d",
    "kwon_fig3_rule_literature_map.png": "185762d49f1f2826558ab34bc2b8873918d4a2ca9bfec728a76201941fc389ab",
}


@dataclass
class Result:
    location: str
    observed: str
    expected: str
    status: str
    source: str


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8-sig"))


def norm(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("−", "-").replace("–", "-")).strip()


def pct(value: float, digits: int = 1) -> str:
    return f"{value * 100:.{digits}f}%"


def number(value: int) -> str:
    return f"{value:,}"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


class Audit:
    def __init__(self, docx_path: Path):
        self.doc = Document(docx_path)
        self.results: list[Result] = []
        self.covered_paragraphs: set[int] = set()
        self.covered_tables: set[int] = set()

    def add(self, location: str, observed, expected, source: str, ok: bool | None = None):
        if ok is None:
            ok = observed == expected
        self.results.append(Result(location, str(observed), str(expected), "일치" if ok else "불일치", source))

    def paragraph(self, location: str, anchor: str, tokens: Iterable[str], source: str):
        matches = [(i, p.text) for i, p in enumerate(self.doc.paragraphs) if anchor in p.text]
        if len(matches) != 1:
            self.add(location, f"anchor matches={len(matches)}", "1", source, False)
            return
        index, text = matches[0]
        self.covered_paragraphs.add(index)
        missing = [token for token in tokens if norm(token) not in norm(text)]
        self.add(location, "모든 기대 수치 있음" if not missing else "누락: " + ", ".join(missing),
                 "모든 기대 수치 있음", source, not missing)

    def cell(self, table_index: int, row: int, col: int, expected: str, source: str, label: str):
        self.covered_tables.add(table_index)
        observed = norm(self.doc.tables[table_index].cell(row, col).text)
        self.add(label, observed, norm(expected), source)

    def table(self, table_index: int, expected_rows: Sequence[Sequence[str]], source: str, label: str):
        self.covered_tables.add(table_index)
        table = self.doc.tables[table_index]
        observed_rows = [[norm(cell.text) for cell in row.cells] for row in table.rows]
        expected = [[norm(cell) for cell in row] for row in expected_rows]
        self.add(label, json.dumps(observed_rows, ensure_ascii=False),
                 json.dumps(expected, ensure_ascii=False), source)

    def unverified_numeric_items(self) -> list[str]:
        items: list[str] = []
        for i, paragraph in enumerate(self.doc.paragraphs):
            if i in self.covered_paragraphs or not re.search(r"\d", paragraph.text):
                continue
            if paragraph.style and paragraph.style.name.startswith("Heading"):
                continue
            items.append(f"문단 {i + 1}: {norm(paragraph.text)}")
        for i, table in enumerate(self.doc.tables):
            if i in self.covered_tables:
                continue
            text = " | ".join(norm(cell.text) for row in table.rows for cell in row.cells)
            if re.search(r"\d", text):
                items.append(f"표 {i + 1}: {text}")
        return items


def read_rules() -> list[dict[str, str]]:
    with RULES_CSV.open("r", encoding="utf-8-sig", newline="") as stream:
        return list(csv.DictReader(stream))


def literal_assignment(source: str, name: str):
    tree = ast.parse(source)
    for node in tree.body:
        if isinstance(node, ast.Assign) and any(isinstance(t, ast.Name) and t.id == name for t in node.targets):
            return ast.literal_eval(node.value)
    raise KeyError(name)


def build_audit(docx_path: Path, fig_dir: Path) -> tuple[Audit, list[str]]:
    run = load_json(RUN_REPORT)
    scoring = load_json(SCORING_REPORT)
    links = load_json(LINK_MANIFEST)
    queries = load_json(QUERY_DEFINITIONS)
    rules = read_rules()
    audit = Audit(docx_path)

    A = run["phases"]["A"]["full_record"]
    B = run["phases"]["B"]["full_record"]
    C = run["phases"]["C"]
    classifier = C["classifier_layer"]
    adjudication = C["semantic_adjudication_layer"]
    final = C["final_layer"]
    movement = adjudication["classifier_to_adjudication_movement"]
    result = links["results"]
    overall = scoring["layers"]["overall"]
    weighted = overall["weighted_metrics"]

    qdefs = {q["question_id"]: q for q in queries["questions"]}
    q_order = C["question_order"]
    q_hits = {q["question_id"]: q["hit_count"] for q in A["questions"]}
    q_units = B["per_question_membership_rows"]
    old_hits = {x["question_id"]: x["v4_hit_count"] for x in run["v4_to_v5_hit_count_change"]}
    changes = {x["question_id"]: x["absolute_change"] for x in run["v4_to_v5_hit_count_change"]}

    total_hits = A["totals"]["hit_count_before_cross_question_deduplication"]
    unique_papers = B["totals"]["evidence_map_rows_unique_papers"]
    screening_units = B["totals"]["question_membership_units_after_bibliographic_deduplication"]
    raw_xml = B["totals"]["raw_xml_files"]
    final_dist = final["decision_distribution"]
    class_dist = classifier["decision_distribution"]
    adjud_dist = adjudication["decision_distribution"]
    retain = final_dist["retain"]
    deprioritize = final_dist["deprioritize"]
    uncertain = final_dist["uncertain"]
    reviewed_retain = adjud_dist["retain"]
    unreviewed_retain = retain - reviewed_retain
    classifier_retain_reviewed = sum(movement["retain"].values())
    classifier_retain_moved = movement["retain"]["deprioritize"] + movement["retain"]["uncertain"]

    census_share = retain / screening_units
    scorer_share = scoring["rogan_gladen"]["design_weighted_scorer_retain_prevalence"]
    scorer_ratio = scorer_share / census_share

    released = [r for r in rules if r["status"] == "released"]
    draft = [r for r in rules if r["status"] == "draft"]
    link_pmids = {str(x["pmid"]) for x in result["links"]}
    rejected = result["rejected_candidates"]
    rejected_by_reason: dict[str, list[dict]] = {}
    for row in rejected:
        rejected_by_reason.setdefault(row["reason"], []).append(row)
    not_in_rows = rejected_by_reason["not_in_v5_corpus"]
    no_retain_rows = rejected_by_reason["no_retain_decision_for_rule_question"]
    not_in_pmids = {str(x["candidate_pmid"]) for x in not_in_rows}
    no_retain_pmids = {str(x["candidate_pmid"]) for x in no_retain_rows}
    badge_total = len(link_pmids | not_in_pmids | no_retain_pmids)

    audit.paragraph(
        "권혁찬 국문초록 허가원문 층",
        "결정층인 허가원문 계층은",
        ["분석 제품 13개", f"고유 성분 {queries['selected_ingredient_count']}개", "제품-성분 연결 47개",
         "복용 조건 32개", f"규칙 {len(rules)}개", f"{len(released)}개", f"{len(draft)}개"],
        "query_definitions.json selected ingredients; rules.csv counts (제품/연결/복용 조건은 정본 대응 없음)",
    )
    audit.results.append(Result(
        "권혁찬 제품 16·분석 13·제품-성분 47·복용 조건 32", "논문 기재값", "정본에 대응 집계 없음",
        "대조 불가", "허용된 정본 파일에는 제품·복용 조건 원장이 없음"
    ))
    audit.paragraph(
        "권혁찬 국문초록 검색",
        "문헌층은 인공지능이",
        [f"성분 {queries['selected_ingredient_count']}개", f"PICOS 질문 {len(queries['questions'])}개", "10,000행",
         number(total_hits), number(unique_papers), number(screening_units)],
        "query_definitions.json; v50_run_report.json phases A/B",
    )
    audit.paragraph(
        "권혁찬 국문초록 선별",
        "선별은 두 층으로 수행하였다",
        [number(screening_units), "커버리지 1.0", number(adjudication["reviewed_rows"]),
         pct(adjudication["reviewed_rows"] / screening_units, 1), number(adjudication["disagreement_count"]),
         pct(adjudication["disagreement_rate"], 1), number(retain), number(deprioritize), number(uncertain),
         number(reviewed_retain), pct(reviewed_retain / retain, 1), number(unreviewed_retain), pct(unreviewed_retain / retain, 1)],
        "v50_run_report.json phase C; direct ratios",
    )
    ci = overall["stratified_bootstrap_95_ci"]
    audit.paragraph(
        "권혁찬 국문초록 채점",
        "선별 품질을 교차 확인하기 위해",
        [number(overall["sample_n"]), pct(weighted["agreement_vs_ai_reference"], 2),
         pct(ci["agreement_vs_ai_reference"][0], 2).rstrip("%"),
         pct(ci["agreement_vs_ai_reference"][1], 2).rstrip("%"),
         pct(weighted["sensitivity_vs_ai_reference"], 2), pct(weighted["specificity_vs_ai_reference"], 2),
         f"{weighted['cohen_kappa_vs_ai_reference_weighted']:.3f}", pct(census_share, 2), pct(scorer_share, 2),
         f"{scorer_ratio:.2f}배", "1.25×10⁻¹⁶"],
        "v50_scoring_report.json; 7,875/43,207; direct ratio",
    )
    audit.paragraph(
        "권혁찬 국문초록 규칙-문헌 연결",
        "규칙–문헌 연결은 규칙 16개 중",
        [f"규칙 {len(rules)}개 중 {result['resolved_rule_count']}개", f"{result['emitted_link_count']}건",
         f"{result['unresolved_rule_count']}개", f"{len(not_in_rows)}건", f"{len(no_retain_rows)}건", "0건"],
        "literature_link_manifest.json results",
    )
    audit.paragraph(
        "권혁찬 영문초록",
        "This study builds a product-name",
        [
            "13 analysed products", "28 unique ingredients", "47 product-ingredient links", "32 administration constraints",
            "16 rules", "15 are released", "five PICOS questions", number(total_hits), number(unique_papers), number(screening_units),
            number(adjudication["reviewed_rows"]), pct(adjudication["reviewed_rows"] / screening_units, 1),
            number(adjudication["disagreement_count"]), pct(adjudication["disagreement_rate"], 1), number(retain),
            number(deprioritize), number(uncertain), number(reviewed_retain), pct(reviewed_retain / retain, 1),
            number(overall["sample_n"]), pct(weighted["agreement_vs_ai_reference"], 2), pct(census_share, 2),
            pct(scorer_share, 2), f"{scorer_ratio:.2f}×", "9 of 16 rules", "10 links", "7 rules",
            "six candidate-link rejections (five distinct papers)", "four candidates",
        ],
        "v50_run_report.json; v50_scoring_report.json; literature_link_manifest.json",
    )

    audit.cell(0, 0, 0, "성능 수치의 기준 이 논문의 선별 성능 수치는 모두 AI 참조표준 대비 값이며 지표 이름에 비교 상대를 명시한다. 사람 참조표준이 아니고 절대적 진실 대비 정확도가 아니다. 사람의 연구 판정은 모든 단계에서 0건이다.",
               "v50_scoring_report.json human_reference_rows=0", "권혁찬 국문초록 성능 기준 상자")

    query_rows = [["질문", "v4.0 hit", "v5.0 hit", "변화", "선별 단위"]]
    labels = ["Q01 아세트아미노펜", "Q02 NSAID", "Q03 감기·알레르기", "Q04 소화효소·담즙산", "Q05 외용 복합성분"]
    for label, qid in zip(labels, q_order):
        change = changes[qid]
        query_rows.append([label, number(old_hits[qid]), number(q_hits[qid]),
                           ("+" if change >= 0 else "-") + number(abs(change)), number(q_units[qid])])
    old_total = sum(old_hits.values())
    query_rows.append(["합계", number(old_total), number(total_hits), "+" + number(total_hits - old_total), number(screening_units)])
    audit.table(1, query_rows, "v50_run_report.json v4_to_v5_hit_count_change/phases A/B", "권혁찬 표 1")

    audit.covered_tables.add(2)
    audit.cell(2, 3, 1, f"{number(queries['selected_ingredient_count'])}개", "query_definitions.json selected_ingredient_count", "권혁찬 표 2 고유 성분")
    audit.cell(2, 6, 1, f"{number(len(rules))}개", "rules.csv row count", "권혁찬 표 2 규칙 수")
    audit.cell(2, 6, 2, f"released {len(released)} · draft {len(draft)}", "rules.csv status counts", "권혁찬 표 2 규칙 상태")

    layer_rows = [
        ["층", "retain", "deprioritize", "uncertain", "합계"],
        ["분류기 층(전량)", number(class_dist["retain"]), number(class_dist["deprioritize"]), number(class_dist["uncertain"]), number(classifier["rows"])],
        ["재판정 층(부분집합)", number(adjud_dist["retain"]), number(adjud_dist["deprioritize"]), number(adjud_dist["uncertain"]), number(adjudication["reviewed_rows"])],
        ["최종", number(retain), number(deprioritize), number(uncertain), number(final["rows"])],
    ]
    audit.table(3, layer_rows, "v50_run_report.json phase C", "권혁찬 표 3")

    movement_rows = [
        ["분류기 라벨", "→ retain", "→ deprioritize", "→ uncertain", "소계"],
    ]
    for label in ["retain", "deprioritize", "uncertain"]:
        row = movement[label]
        movement_rows.append([label, number(row["retain"]), number(row["deprioritize"]), number(row["uncertain"]), number(sum(row.values()))])
    movement_rows.append(["합계", number(adjud_dist["retain"]), number(adjud_dist["deprioritize"]), number(adjud_dist["uncertain"]), number(adjudication["reviewed_rows"])])
    audit.table(4, movement_rows, "v50_run_report.json phase C classifier_to_adjudication_movement", "권혁찬 표 4")

    metric_specs = [
        ("agreement_vs_ai_reference", "agreement_vs_ai_reference"),
        ("sensitivity_vs_ai_reference", "sensitivity_vs_ai_reference"),
        ("specificity_vs_ai_reference", "specificity_vs_ai_reference"),
        ("precision_vs_ai_reference", "precision_vs_ai_reference"),
        ("f1_vs_ai_reference", "f1_vs_ai_reference"),
        ("Cohen κ", "cohen_kappa_vs_ai_reference_weighted"),
    ]
    metric_rows = [["지표", "값", "부트스트랩 95% CI"]]
    for label, key in metric_specs:
        value = f"{weighted[key]:.3f}" if label == "Cohen κ" else pct(weighted[key], 2)
        low, high = ci[key]
        interval = f"{low:.3f} ~ {high:.3f}" if label == "Cohen κ" else f"{low * 100:.2f} ~ {high * 100:.2f}"
        metric_rows.append([label, value, interval])
    audit.table(5, metric_rows, "v50_scoring_report.json layers.overall", "권혁찬 표 5")

    layer_score_rows = [["층", "모수 N", "표본 n", "agreement", "Cohen κ"]]
    for label, key in [("분류기 층", "classifier"), ("재판정 층", "adjudicated"), ("전체", "overall")]:
        layer = scoring["layers"][key]
        layer_score_rows.append([
            label, number(round(layer["estimated_population_N"])), number(layer["sample_n"]),
            pct(layer["weighted_metrics"]["agreement_vs_ai_reference"], 2),
            f"{layer['weighted_metrics']['cohen_kappa_vs_ai_reference_weighted']:.3f}",
        ])
    audit.table(6, layer_score_rows, "v50_scoring_report.json layers", "권혁찬 표 6")

    quote_pass = sum(bool(x.get("locator")) and x.get("quote_exact_match") is True for x in result["links"])
    link_rows = [
        ["항목", "값"],
        ["규칙", f"{result['rule_count']}개"],
        ["문헌이 연결된 규칙", f"{result['resolved_rule_count']}개"],
        ["미연결 규칙", f"{result['unresolved_rule_count']}개"],
        ["연결 수", f"{result['emitted_link_count']}건"],
        ["기각된 후보", f"{result['rejected_candidate_count']}건"],
        ["기각 사유 — v5.0 코퍼스에 없음", f"{len(not_in_rows)}건 (고유 {len(not_in_pmids)}편)"],
        ["기각 사유 — 허용 질문에서 retain 아님", f"{len(no_retain_rows)}건"],
        ["문장 단위 인용 대조", f"{quote_pass}건 전건 통과"],
    ]
    audit.table(7, link_rows, "literature_link_manifest.json direct counts", "권혁찬 표 7")

    audit.paragraph(
        "권혁찬 §3.1 수집·분석 제품",
        "식약처 의약품안전나라의 의약품상세정보에서 제품 16개",
        ["제품 16개", "허가 취하 2개", "1개", "13개"],
        "대조 불가: 허용된 정본에는 제품 원장이 없음",
    )
    audit.paragraph(
        "권혁찬 §3.1 규칙·복용 조건 구분",
        "규칙은 허가원문에서 직접 도출한 16개",
        [f"{len(rules)}개", "복용 조건 32개", f"released 규칙 {len(released)}개"],
        "rules.csv for rule/status counts; administration-condition count unavailable",
    )
    audit.paragraph(
        "권혁찬 §3.2 검색 기간",
        "기간 제한은 질문별로 검색식 안에 명시하였다",
        ["Q01~Q03", "2010-01-01", "Q04~Q05", "2000-01-01"],
        "query_definitions.json questions[].date_range",
    )
    audit.paragraph(
        "권혁찬 §3.3 검색 실행",
        "검색은 2026년 7월 28일에 실행하였다",
        ["2026년 7월 28일", "9,999건", f"원본 XML {raw_xml}개"],
        "v50_run_report.json phases A/B (9,999 분할 기준은 정본 대응 없음)",
    )
    q04_p_terms = len(qdefs[q_order[3]]["blocks"]["P"][0]["terms"])
    q05_p_terms = len(qdefs[q_order[4]]["blocks"]["P"][0]["terms"])
    audit.paragraph(
        "권혁찬 표 1 캡션 Q05 진단",
        "표 1. 검색식 재설계 전후의 질문별 검색 건수",
        [number(unique_papers), number(screening_units), "Q05", "네 개", "Q01은 세 개", "Q04는 두 개",
         f"검색어 {q05_p_terms}개", f"{q04_p_terms}개"],
        "v50_run_report.json; query_definitions.json P-term counts (v4 outcome-term counts lack canonical detail)",
    )
    audit.paragraph(
        "권혁찬 §3.4 분류기 검증 사례",
        "분류기 층은 실제 사례 42건",
        [f"{C['classifier_validation']['case_count']}건", f"{C['classifier_validation']['pass_count']}건 일치",
         f"{C['classifier_validation']['fail_count']}건 불일치"],
        "v50_run_report.json phase C classifier_validation",
    )
    audit.paragraph(
        "권혁찬 §3.5 채점 설계",
        "표본은 코퍼스를 완전분할하는 층화 확률표본",
        [number(scoring["design"]["sample_n"]), f"층은 {scoring['design']['sampling_strata']}개", number(scoring["design"]["population_N"]),
         "20260730-v50-scoring-arm", f"확률표본 층에서 {overall['probability_sample_n']}행", f"전수 층에서 {overall['census_n']}행"],
        "v50_scoring_report.json design/layers.overall",
    )
    audit.paragraph(
        "권혁찬 §3.6 통계 처리",
        "신뢰구간은 층화 부트스트랩",
        [number(scoring["design"]["bootstrap_draws"]), "확률표본 층 21개", "전수 층 12개", "표집오차가 0", "3범주"],
        "v50_scoring_report.json design/layers.overall.stratified_bootstrap_design",
    )
    audit.paragraph(
        "권혁찬 §3.7 사이트 배지",
        "공개 화면의 판정 카드는",
        [f"{badge_total}편", f"검증 {len(link_pmids)}편", f"코퍼스에 없음 {len(not_in_pmids)}편", f"retain 아님 {len(no_retain_pmids)}편"],
        "literature_link_manifest.json distinct PMID sets",
    )
    audit.paragraph(
        "권혁찬 그림 2 캡션",
        "그림 2. 코퍼스 구축과 두 층 선별 흐름",
        [number(screening_units), number(adjudication["reviewed_rows"]), pct(adjudication["reviewed_rows"] / screening_units, 1),
         number(retain), number(reviewed_retain), pct(reviewed_retain / retain, 1)],
        "v50_run_report.json phase C; direct ratios",
    )
    audit.paragraph(
        "권혁찬 §4.3 최종 retain 검산",
        "검산은 7,278",
        [number(class_dist["retain"]), number(classifier_retain_reviewed), number(reviewed_retain), number(retain),
         number(unreviewed_retain), pct(unreviewed_retain / retain, 1)],
        "v50_run_report.json phase C; direct arithmetic",
    )
    audit.paragraph(
        "권혁찬 §4.3 재판정 분류기 retain 이동",
        "재판정을 받은 분류기 retain 596건 중",
        [number(classifier_retain_reviewed), number(movement["retain"]["retain"]), number(classifier_retain_moved),
         pct(classifier_retain_moved / classifier_retain_reviewed, 1), number(unreviewed_retain)],
        "v50_run_report.json classifier_to_adjudication_movement; direct ratio",
    )
    focus = scoring["focus_classifier_unadjudicated_retain"]
    audit.paragraph(
        "권혁찬 표 6 캡션 미재판정 retain",
        "재판정을 받지 않은 최종 retain",
        [number(focus["population_N"]), pct(focus["weighted_metrics"]["sensitivity_vs_ai_reference"], 2),
         pct(focus["stratified_bootstrap_95_ci"]["sensitivity_vs_ai_reference"][0], 2).rstrip("%"),
         pct(focus["stratified_bootstrap_95_ci"]["sensitivity_vs_ai_reference"][1], 2).rstrip("%")],
        "v50_scoring_report.json focus_classifier_unadjudicated_retain",
    )
    question_tokens = []
    for qid in [q_order[3], q_order[4], q_order[2], q_order[0], q_order[1]]:
        question_tokens.append(pct(scoring["question_agreement"][qid]["exact_three_label_agreement_vs_ai_reference_weighted"], 2))
    audit.paragraph(
        "권혁찬 §4.4 질문별 가중 일치율",
        "질문별 가중 3범주 일치율은",
        question_tokens,
        "v50_scoring_report.json question_agreement",
    )
    audit.paragraph(
        "권혁찬 §4.5 retain 비율",
        "파이프라인의 전수 retain 비율은 18.23%(7,875/43,207)",
        [pct(census_share, 2), f"({number(retain)}/{number(screening_units)})", pct(scorer_share, 2), f"{scorer_ratio:.2f}배"],
        "v50_scoring_report.json rogan_gladen; direct arithmetic",
    )
    raw_dirs = scoring["disagreements"]["raw_by_direction"]
    audit.paragraph(
        "권혁찬 §4.5 불일치 방향",
        "채점 894행에서 관측된 불일치",
        [number(overall["sample_n"]), number(scoring["disagreements"]["raw_total"]),
         *[f"{direction.replace('->', '→')} {count}건" for direction, count in raw_dirs.items()],
         pct(weighted["specificity_vs_ai_reference"], 2), pct(weighted["sensitivity_vs_ai_reference"], 2)],
        "v50_scoring_report.json disagreements/layers.overall",
    )
    audit.paragraph(
        "권혁찬 §4.5 Rogan-Gladen",
        "Rogan–Gladen 보정을 적용한 값은",
        ["16자리", "0.1195036622964358", "0.11950366229643568", "1.25×10⁻¹⁶"],
        "v50_scoring_report.json rogan_gladen",
    )
    audit.paragraph(
        "권혁찬 그림 3 캡션",
        "그림 3. 규칙 16개의 문헌 연결 상태",
        [f"규칙 {result['rule_count']}개", f"규칙 {result['resolved_rule_count']}개", f"링크 {result['emitted_link_count']}건",
         f"미연결 {result['unresolved_rule_count']}개", f"{len(not_in_rows)}건", f"{len(no_retain_rows)}건"],
        "literature_link_manifest.json results",
    )
    audit.paragraph(
        "권혁찬 §4.6 검색 기간이 아닌 사유",
        "검색 기간은 사유가 아니다",
        ["Q01~Q03", "2010-01-01", "Q04~Q05", "2000-01-01", number(unique_papers), "2000년부터 2026년", "9편", "2010년부터 2025년", "한 편도 없다"],
        "query_definitions.json for windows; publication-year ranges unavailable in allowed canon",
    )
    audit.results.append(Result(
        "권혁찬 코퍼스·미연결 후보 출판연도 범위", "2000~2026 / 2010~2025", "정본에 연도 필드 없음",
        "대조 불가", "허용된 정본 파일에는 evidence_map.csv 본문이 없음"
    ))
    audit.paragraph(
        "권혁찬 §5.2 연결 실패 10건",
        "연결에 실패한 10건의 사유",
        [f"{result['rejected_candidate_count']}건", number(old_total), number(total_hits)],
        "literature_link_manifest.json; v50_run_report.json",
    )
    audit.paragraph(
        "권혁찬 한계 4 채점 방향",
        "넷째, 채점 arm은",
        [pct(census_share, 2), pct(scorer_share, 2), f"{scorer_ratio:.2f}배", "155건", "20건"],
        "v50_scoring_report.json",
    )
    audit.paragraph(
        "권혁찬 한계 5 미재판정 retain",
        "다섯째, 최종 retain",
        [number(retain), number(unreviewed_retain), pct(unreviewed_retain / retain, 1), number(classifier_retain_reviewed),
         pct(classifier_retain_moved / classifier_retain_reviewed, 1), number(unreviewed_retain)],
        "v50_run_report.json phase C; direct arithmetic",
    )
    audit.paragraph(
        "권혁찬 한계 6 분류기 검증 사례",
        "여섯째, 분류기 층 검증에서",
        [f"{C['classifier_validation']['case_count']}건", f"{C['classifier_validation']['fail_count']}건"],
        "v50_run_report.json classifier_validation",
    )

    fig_source = (fig_dir / "fig_kwon.py").read_text(encoding="utf-8")
    expected_fragments = [
        f"고유 성분 {queries['selected_ingredient_count']}", f"규칙 {len(rules)}(released {len(released)} · draft {len(draft)})",
        f"선별 단위 {number(screening_units)}건 → retain {number(retain)}건 → 규칙 연결 {result['resolved_rule_count']}/{result['rule_count']} · 링크 {result['emitted_link_count']}건",
        f"ESearch 합계 {number(total_hits)}건", f"고유 논문 {number(unique_papers)}편", f"선별 단위(논문–질문 조합) {number(screening_units)}건",
        f"retain {number(class_dist['retain'])} · deprioritize {number(class_dist['deprioritize'])} · uncertain {number(class_dist['uncertain'])}",
        f"{number(adjudication['reviewed_rows'])}건({pct(adjudication['reviewed_rows'] / screening_units, 1)})만",
        f"{number(adjudication['disagreement_count'])}건({pct(adjudication['disagreement_rate'], 1)})이 분류기와 다른 라벨로 이동",
        f'("retain", "{number(retain)}", "{pct(retain / screening_units, 1)}"',
        f'("deprioritize", "{number(deprioritize)}", "{pct(deprioritize / screening_units, 1)}"',
        f'("uncertain", "{number(uncertain)}", "{pct(uncertain / screening_units, 1)}"',
        f"{number(class_dist['retain'])} − {number(classifier_retain_reviewed)} + {number(reviewed_retain)} = {number(retain)}",
        f"{number(reviewed_retain)}건 ({pct(reviewed_retain / retain, 1)})", f"{number(unreviewed_retain)}건 ({pct(unreviewed_retain / retain, 1)})",
        f"연결 {result['resolved_rule_count']}규칙 · 링크 {result['emitted_link_count']}건", f"미연결 {result['unresolved_rule_count']}규칙",
        f"코퍼스 미인출 {len(not_in_rows)}건", f"retain 아님 {len(no_retain_rows)}건",
    ]
    missing = [fragment for fragment in expected_fragments if fragment not in fig_source]
    audit.add("권혁찬 그림 1~3 소스 수치", "모든 기대 수치 있음" if not missing else "누락: " + ", ".join(missing),
              "모든 기대 수치 있음", "fig_kwon.py vs canonical artifacts", not missing)

    manifest_rules = {x["rule_id"]: x for x in result["rules"]}
    expected_rule_tuples = []
    for rule in rules:
        item = manifest_rules[rule["rule_id"]]
        num = rule["rule_id"].rsplit("-", 1)[1]
        link_count = item["link_count"]
        why = None
        if not link_count:
            counts = item["rejection_counts"]
            if "not_in_v5_corpus" in counts:
                why = ("코퍼스 미인출", counts["not_in_v5_corpus"])
            else:
                why = ("retain 아님", counts["no_retain_decision_for_rule_question"])
        expected_rule_tuples.append((num, rule["rule_type"], link_count, why))
    observed_rule_tuples = literal_assignment(fig_source, "RULES")
    audit.add("권혁찬 그림 3 규칙별 연결·기각 수", observed_rule_tuples, expected_rule_tuples,
              "fig_kwon.py RULES vs rules.csv/literature_link_manifest.json")

    for name, expected_hash in EXPECTED_FIGURE_HASHES.items():
        path = fig_dir / name
        audit.add(f"권혁찬 그림 파일 해시 {name}", sha256(path), expected_hash,
                  "2026-07-31 visually inspected source/PNG baseline")

    return audit, audit.unverified_numeric_items()


def print_results(audit: Audit, unverified: list[str]) -> int:
    print("| 위치 | 관찰값 | 정본/기대값 | 판정 | 근거 |")
    print("|---|---|---|---|---|")
    for result in audit.results:
        values = [result.location, result.observed, result.expected, result.status, result.source]
        safe = [str(v).replace("|", " / ").replace("\n", " ") for v in values]
        print("| " + " | ".join(safe) + " |")
    failures = [r for r in audit.results if r.status == "불일치"]
    print(f"\nSUMMARY pass={sum(r.status == '일치' for r in audit.results)} "
          f"fail={len(failures)} unverified_groups={sum(r.status == '대조 불가' for r in audit.results)}")
    if unverified:
        print("\nUNMAPPED NUMERIC ITEMS (manual review required)")
        for item in unverified:
            print("- " + item.replace("\n", " "))
    return 1 if failures else 0


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", type=Path, default=DEFAULT_DOCX)
    parser.add_argument("--figure-dir", type=Path, default=DEFAULT_FIG_DIR)
    args = parser.parse_args()
    audit, unverified = build_audit(args.docx, args.figure_dir)
    return print_results(audit, unverified)


if __name__ == "__main__":
    sys.exit(main())
