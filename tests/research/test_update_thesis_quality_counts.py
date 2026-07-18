from pathlib import Path

from docx import Document

from scripts.research.update_thesis_quality_counts import update


def test_updates_three_quality_count_mentions(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("연구 테스트 31건")
    doc.add_paragraph("31 research tests")
    doc.add_paragraph("연구 테스트 31건")
    doc.save(source)
    assert update(source, output) == 3
    text = "\n".join(p.text for p in Document(output).paragraphs)
    assert "31건" not in text
    assert "31 research tests" not in text
    assert text.count("35") == 3
