from __future__ import annotations

import csv
import hashlib
import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "research_v3" / "human_review_minimal"


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def write_csv(path: Path, rows: list[dict[str, str]], fields: list[str] | None = None) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fields = fields or list(rows[0])
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def font(run, family="Pretendard", size=10.5, bold=False, color="222222"):
    run.font.name = family
    run._element.rPr.rFonts.set(qn("w:eastAsia"), family)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_title(doc: Document, text: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(5)
    font(p.add_run(text), "Pretendard ExtraBold", 23, False, "123A63")
    if subtitle:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(14)
        font(p.add_run(subtitle), "Pretendard Medium", 10.5, False, "4E6477")


def add_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(11)
    p.paragraph_format.space_after = Pt(5)
    font(p.add_run(text), "Pretendard ExtraBold", 16, False, "123A63")


def add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    font(p.add_run(text), "Pretendard SemiBold", 12, False, "1769AA")


def add_p(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix and text.startswith(bold_prefix):
        font(p.add_run(bold_prefix), "Pretendard SemiBold", 10.5, False, "222222")
        font(p.add_run(text[len(bold_prefix):]))
    else:
        font(p.add_run(text))


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.55)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(3)
        font(p.add_run(item))


def add_callout(doc: Document, title: str, body: str, fill="EAF4FB") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(17.4)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    font(p.add_run(title), "Pretendard SemiBold", 10.5, False, "123A63")
    p = cell.add_paragraph()
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(2)
    font(p.add_run(body), "Pretendard", 10, False, "243746")


def add_task_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    widths = [Cm(3.1), Cm(9.8), Cm(4.5)]
    for idx, width in enumerate(widths):
        table.columns[idx].width = width
    headers = ["작업", "확인할 것", "완료 표시"]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "123A63")
        p = cell.paragraphs[0]
        font(p.add_run(text), "Pretendard SemiBold", 9.5, False, "FFFFFF")
    set_repeat_table_header(table.rows[0])
    for a, b, c in rows:
        cells = table.add_row().cells
        for i, text in enumerate((a, b, c)):
            if len(table.rows) % 2 == 1:
                set_cell_shading(cells[i], "F5F8FA")
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            font(p.add_run(text), "Pretendard", 9.2, False, "222222")


def build_docx(path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    normal = doc.styles["Normal"]
    normal.font.name = "Pretendard"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Pretendard")
    normal.font.size = Pt(10.5)

    add_title(doc, "사람이 확인할 것만 남긴 연구 종료 가이드", "권혁찬 졸업논문 연구 · Codex AI 준비본 · 2026-07-13")
    add_callout(doc, "먼저 결론", "기술 구현·AI 검토·초안·패킷 구성은 끝났다. 사람은 아래 5묶음의 빈칸만 확인한다. 실제 확인 전에는 ‘사람 검토 완료’, ‘전문가 승인’, ‘체계적 문헌고찰 완료’라고 쓰지 않는다.")
    add_h1(doc, "1. 무엇을 하면 끝나는가")
    add_task_table(doc, [
        ("① 연구 승인", "연구 방향, 제목, 학번, 제출 허용", "승인서 1장"),
        ("② 검색식", "5개 검색식 × PRESS 7항목", "35칸 판정"),
        ("③ 우선 문헌", "AI가 추린 118건의 제목·초록", "포함/제외/불확실"),
        ("④ 안전 규칙", "6개 규칙의 수치·범위·문구·locator", "승인/수정/반려"),
        ("⑤ 독립 평가", "12개 입력 시나리오의 정답 위험", "정답 잠금"),
    ])
    add_p(doc, "권장 순서: ① → ② → ③ → ④ → ⑤. 각 CSV의 흰 빈칸만 입력하면 된다.")

    doc.add_page_break()
    add_title(doc, "읽기 전에 알아둘 것")
    add_h1(doc, "2. 이미 끝난 것")
    add_bullets(doc, [
        "연구 방향: 고함량 영양성분·함량 중심(비타민 D·칼슘, B6, 철, 마그네슘, 아연).",
        "문헌 검색 원본·중복 제거·전체 큐 15,890건 보존.",
        "Codex AI 1차 검토: 포함 후보 6,106건, 불확실 9,504건, 제외 후보 280건.",
        "우선 문헌 118건, 검색전략 5개, 결정론적 규칙 초안 6개 준비.",
        "연구 테스트 20/20, 앱 테스트 39/39, lint·typecheck·build 통과.",
        "논문 evidence-bound 초안 DOCX/PDF 생성 및 렌더 검수.",
    ])
    add_h1(doc, "3. 완료 주장을 어디까지 할 수 있는가")
    add_callout(doc, "최소 부담 경로", "118건을 사람이 확인하면 ‘Codex AI 보조 근거지도와 우선 근거 검토’로 제출 가능하다. 이는 15,890건 전수 사람이 판정한 체계적 문헌고찰과 같지 않다.", "E8F5EE")
    add_callout(doc, "전수 경로", "‘체계적 문헌고찰 완료’라고 쓰려면 15,890건 전체 제목·초록 판정, 전문 입수·판정, 추출 검증이 추가로 필요하다.", "FFF4E5")

    doc.add_page_break()
    add_title(doc, "검토 ① 연구 승인")
    add_h1(doc, "4. 승인자는 네 가지만 본다")
    add_bullets(doc, [
        "연구 방향 유지: 고함량 영양성분·함량 중심.",
        "연구자 정보: 권혁찬 / 2021194024 / 연세대학교 약학대학.",
        "논문 제목·주장 수준: AI 보조 근거지도 및 안전성 정보 시스템.",
        "검토 완료 후 canonical 최종본 승격 허용 여부.",
    ])
    add_p(doc, "입력 파일: 01_연구_제출_승인서.csv")
    add_callout(doc, "완료 기준", "결정값, 이름, 역할, 날짜가 모두 채워지고 required_revision이 비어 있거나 반영 완료 상태여야 한다.")

    add_h1(doc, "5. 승인자가 읽을 핵심 문장")
    add_p(doc, "본 연구는 고함량 영양성분의 함량·상한섭취량·조건별 안전성 정보를 근거 위치와 함께 구조화하고, 사용자가 입력한 섭취량에서 확인이 필요한 상황을 안내하는 연구용 시스템을 개발했다. 진단·처방·개인 맞춤 의료판정을 제공하지 않는다.")

    doc.add_page_break()
    add_title(doc, "검토 ② 검색식 · ③ 우선 문헌")
    add_h1(doc, "6. 검색식 5개")
    add_p(doc, "입력 파일: 02_PRESS_검색전략_검토.csv")
    add_bullets(doc, [
        "rating: yes / no / unclear 중 하나.",
        "no 또는 unclear이면 comment와 required_change 작성.",
        "수정 반영 뒤 resolution과 status=completed 입력.",
        "질문은 연구질문 일치, 용어, 구문, 제한, 번역, 구조, 재현성 7항목.",
    ])
    add_h1(doc, "7. 우선 문헌 118건")
    add_p(doc, "입력 파일: 03_우선문헌_118건_검토.csv")
    add_bullets(doc, [
        "제목·초록 읽고 include_candidate / exclude / uncertain 중 하나.",
        "exclude이면 정해진 제외 이유 코드 하나 기록.",
        "include_candidate 또는 uncertain이면 전문 입수 대상으로 보냄.",
        "전문 미확보는 exclude가 아니다. retrieval 상태로 기록.",
        "AI 제안은 우선순위 참고값. 그대로 복사하지 않음.",
    ])
    add_callout(doc, "가장 중요한 원칙", "논문 한 편의 결론을 전체 영양성분 규칙으로 확대하지 않는다. 대상·용량·기간·결과·한계를 함께 확인한다.")

    doc.add_page_break()
    add_title(doc, "검토 ④ 규칙 · ⑤ 독립 평가")
    add_h1(doc, "8. 안전 규칙 6개")
    add_p(doc, "입력 파일: 04_규칙_6건_검토.csv")
    add_bullets(doc, [
        "threshold_correct: 수치·단위가 원문과 같은가.",
        "scope_correct: 총섭취량/보충제/식이 범위가 맞는가.",
        "conditions_correct·exceptions_correct: 적용조건과 예외가 안전한가.",
        "message_safe·next_action_safe: 과장 없이 다음 행동을 안내하는가.",
        "source_locator_verified: 지정 페이지·표·행에서 직접 확인했는가.",
        "overall_decision: approve / revise / reject.",
    ])
    add_h1(doc, "9. 독립 시나리오 12개")
    add_p(doc, "입력 파일: 05_독립시나리오_12건_확정.csv")
    add_bullets(doc, [
        "input_json은 개발 시나리오에서 복사한 입력 틀. 정답은 비워 둠.",
        "평가자는 gold_hazards_json을 독립 작성.",
        "정답 작성 후 locked_before_test=true로 잠금.",
        "그 뒤 시스템 예측을 실행하고 민감도·중대한 위음성·정확 일치율 계산.",
    ])
    add_callout(doc, "독립성", "개발자가 만든 expected_hazards_json은 정답 칸에 사전입력하지 않았다. 평가자는 규칙 근거를 보고 독립 판정한다.")

    doc.add_page_break()
    add_title(doc, "파일 위치와 종료 조건")
    add_h1(doc, "10. 작업 파일")
    add_bullets(doc, [
        "00_읽기_가이드.md: 이 문서의 텍스트판.",
        "01_연구_제출_승인서.csv: 연구·신원·제출 승인.",
        "02_PRESS_검색전략_검토.csv: 35개 판정.",
        "03_우선문헌_118건_검토.csv: 제목·초록 우선 검토.",
        "04_규칙_6건_검토.csv: 규칙·근거 위치 확인.",
        "05_독립시나리오_12건_확정.csv: 독립 정답 작성.",
        "manifest.json: 파일 수·행 수·SHA-256·빈칸 사전입력 검사.",
    ])
    add_h1(doc, "11. 연구 종료 체크")
    add_task_table(doc, [
        ("기술", "테스트·빌드·패키지 검사 통과", "완료"),
        ("승인", "연구 방향·학번·제출 승인", "미완료"),
        ("검색", "PRESS 35개 항목 완료", "미완료"),
        ("문헌", "우선 118건 + 필요한 전문 검토", "미완료"),
        ("규칙", "6개 규칙 승인 후 released", "미완료"),
        ("평가", "독립 12개 정답 잠금·평가", "미완료"),
        ("제출", "논문 canonical 승격·최종 PDF 확인", "미완료"),
    ])
    add_callout(doc, "현재 상태", "release_ready=false. 위 미완료 항목이 실제로 채워지면 verifier를 다시 실행하고 최종본을 승격한다.", "FFF4E5")

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font(footer.add_run("Codex AI 준비본 · 실제 확인 전 사람 검토 완료로 간주하지 않음"), "Pretendard", 8, False, "657786")
    doc.save(path)


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    press = read_csv(ROOT / "research_v3/search/provisional_pubmed_20260710/peer_review.csv")
    priority = read_csv(ROOT / "research_v3/screening/review_packets/priority_118_review_packet.csv")
    rules = read_csv(ROOT / "research_v3/rules/expert_rule_review_packet.csv")
    dev = read_csv(ROOT / "research_v3/validation/development_scenarios.csv")

    approvals = [
        {"item_id": "APP-01", "question": "고함량 영양성분·함량 중심 연구 방향을 승인하는가", "decision": "", "required_revision": "", "reviewer_id": "", "reviewer_role": "", "reviewed_at": ""},
        {"item_id": "APP-02", "question": "권혁찬 학번 2021194024가 학교 기록과 일치하는가", "decision": "", "required_revision": "", "reviewer_id": "", "reviewer_role": "", "reviewed_at": ""},
        {"item_id": "APP-03", "question": "AI 보조 근거지도 및 안전성 정보 시스템이라는 주장 수준을 승인하는가", "decision": "", "required_revision": "", "reviewer_id": "", "reviewer_role": "", "reviewed_at": ""},
        {"item_id": "APP-04", "question": "모든 검토 완료 후 canonical 최종본 승격과 제출을 허용하는가", "decision": "", "required_revision": "", "reviewer_id": "", "reviewer_role": "", "reviewed_at": ""},
    ]
    write_csv(OUT / "01_연구_제출_승인서.csv", approvals)
    shutil.copy2(ROOT / "research_v3/search/provisional_pubmed_20260710/peer_review.csv", OUT / "02_PRESS_검색전략_검토.csv")
    shutil.copy2(ROOT / "research_v3/screening/review_packets/priority_118_review_packet.csv", OUT / "03_우선문헌_118건_검토.csv")
    shutil.copy2(ROOT / "research_v3/rules/expert_rule_review_packet.csv", OUT / "04_규칙_6건_검토.csv")

    independent = []
    for idx, row in enumerate(dev, start=1):
        independent.append({
            "scenario_id": f"IND-{idx:03d}",
            "scenario_type": row["scenario_type"],
            "input_json": row["input_json"],
            "gold_hazards_json": "",
            "critical": row["critical"],
            "adjudicator_id": "",
            "adjudicated_at": "",
            "locked_before_test": "false",
            "source_note": "input scaffold copied from development scenario; gold label intentionally blank",
            "notes": "",
        })
    write_csv(OUT / "05_독립시나리오_12건_확정.csv", independent)

    guide = """# 사람이 확인할 것만 남긴 연구 종료 가이드

## 결론

기술 구현, Codex AI 검토, 초안, 패킷 구성은 끝났다. 사람은 다음 5묶음의 빈칸만 확인한다.

1. 연구 방향·학번·제출 승인 4건
2. 검색전략 PRESS 35건
3. 우선 문헌 제목·초록 118건과 필요한 전문
4. 안전 규칙 6건
5. 독립 평가 시나리오 12건

## 주장 범위

- 최소 부담 경로: 118건 확인 후 `Codex AI 보조 근거지도와 우선 근거 검토`로 보고한다.
- 전수 경로: `체계적 문헌고찰 완료` 주장은 15,890건 전체 제목·초록 판정과 전문 검토가 필요하다.
- Codex AI 결과를 사람 또는 전문가 검토로 표시하지 않는다.

## 입력 순서

`01_연구_제출_승인서.csv` → `02_PRESS_검색전략_검토.csv` → `03_우선문헌_118건_검토.csv` → `04_규칙_6건_검토.csv` → `05_독립시나리오_12건_확정.csv`

## 종료 조건

모든 필수 판정·검토자·검토일·근거 위치가 채워지고, 독립 정답이 시험 전에 잠기고, verifier가 통과한 뒤 `release_ready=true`와 canonical 최종본 승격을 별도 수행한다.
"""
    (OUT / "00_읽기_가이드.md").write_text(guide, encoding="utf-8")
    build_docx(OUT / "사람_최소검토_가이드.docx")

    tracked = [p for p in OUT.iterdir() if p.is_file() and p.name != "manifest.json"]
    manifest = {
        "generated_at": "2026-07-13",
        "generator": "Codex AI",
        "release_ready": False,
        "claim_mode": "ai_assisted_evidence_map_priority_review",
        "counts": {"approval_items": 4, "press_items": len(press), "priority_records": len(priority), "rule_items": len(rules), "independent_scenarios": len(independent)},
        "human_prefill_counts": {"approval_decisions": 0, "press_ratings": 0, "priority_decisions": 0, "rule_decisions": 0, "independent_gold_labels": 0},
        "files": [{"name": p.name, "bytes": p.stat().st_size, "sha256": sha256(p)} for p in sorted(tracked)],
    }
    (OUT / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
