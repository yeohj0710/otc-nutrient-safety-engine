from __future__ import annotations

import argparse
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


FONT = "Pretendard"
NAVY = "17233C"
LIGHT_BLUE = "E8EEF7"
WARNING = "FFF3CD"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margin(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_font(run, size=10.5, bold=False, color=None) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_page_number(paragraph) -> None:
    for child in list(paragraph._p):
        paragraph._p.remove(child)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("- ")
    set_font(run, 9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, end):
        run._r.append(node)
    set_font(paragraph.add_run(" -"), 9)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.65
    normal.paragraph_format.space_after = Pt(7)
    for style_name, size, before, after in (
        ("Title", 20, 0, 18),
        ("Heading 1", 15, 18, 8),
        ("Heading 2", 12.5, 14, 6),
        ("Heading 3", 11, 10, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(NAVY)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def clear_body(doc: Document) -> None:
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def parse_markdown(path: Path):
    lines = path.read_text(encoding="utf-8").splitlines()
    title = lines[0].lstrip("\ufeff").removeprefix("# ").strip()
    sections: dict[str, list[tuple[int, str]]] = {}
    current = "meta"
    sections[current] = []
    for raw in lines[1:]:
        line = raw.rstrip()
        if line.startswith("## "):
            current = line[3:].strip()
            sections[current] = []
        elif line.startswith("### "):
            sections[current].append((3, line[4:].strip()))
        elif line.startswith("> "):
            sections[current].append((9, line[2:].strip()))
        elif re.match(r"^\d+\. ", line) and current == "참고문헌 초안":
            sections[current].append((8, line))
        elif line.startswith("주요어:") or line.startswith("Keywords:"):
            sections[current].append((7, line))
        elif line.startswith("Reference basis for Korean prose rhythm:"):
            continue
        elif line.strip():
            sections[current].append((0, line.strip().replace("  ", "")))
    return title, sections


def add_text_block(doc: Document, items, body_size: float = 10.5, keyword_size: float = 9.5) -> None:
    for kind, text in items:
        if kind == 3:
            doc.add_heading(text, level=2)
        elif kind == 9:
            table = doc.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            shade(cell, WARNING)
            set_cell_margin(cell, 140, 160, 140, 160)
            p = cell.paragraphs[0]
            r = p.add_run("연구 상태: " + text)
            set_font(r, 9.5, True, "7A4E00")
        elif kind == 7:
            p = doc.add_paragraph()
            r = p.add_run(text)
            set_font(r, keyword_size, True, NAVY)
        elif kind == 8:
            p = doc.add_paragraph(style="List Number")
            r = p.add_run(re.sub(r"^\d+\.\s*", "", text))
            set_font(r, 9.5)
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.first_line_indent = Cm(-0.6)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.7)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_font(p.add_run(text), body_size)


def add_cover(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(55)
    set_font(p.add_run("졸 업 논 문"), 22, True, NAVY)
    table = doc.add_table(rows=9, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(4.1)
    table.columns[1].width = Cm(11.4)
    rows = [
        ("연구 주제명", f"국문: {title}\n영문: Safety Assessment of High-Dose Nutrient Intake Standards and Development of a Personalized Query Tool"),
        ("성명", "권혁찬"), ("학번", "2021194024"),
        ("실습기간", "2026년 3월 3일 - 2026년 6월 19일"),
        ("실습장소", "연세대학교 약학대학"), ("논문양식", "종설논문 / 방법론 개발 연구"),
        ("담당교수", "장민정 교수님 (전자 승인 2026-07-13)"), ("제출일", "2026년 7월 13일"),
        ("문서상태", "evidence-bound 작업본 · canonical 최종본 아님"),
    ]
    for i, (label, value) in enumerate(rows):
        left, right = table.rows[i].cells
        left.width, right.width = Cm(4.1), Cm(11.4)
        shade(left, LIGHT_BLUE)
        for c in (left, right):
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margin(c)
        set_font(left.paragraphs[0].add_run(label), 9.5, True, NAVY)
        for j, part in enumerate(value.split("\n")):
            if j:
                right.paragraphs[0].add_run().add_break()
            set_font(right.paragraphs[0].add_run(part), 9.3)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    set_font(p.add_run("연세대학교 약학대학 전공심화실습(1)"), 13, True, NAVY)
    doc.add_page_break()


def add_contents(doc: Document) -> None:
    heading = doc.add_heading("목차", level=1)
    heading.paragraph_format.page_break_before = True
    entries = [
        "국문초록", "영문초록", "1. 서론", "2. 연구 방법", "3. 연구 결과",
        "4. 고찰", "5. 결론", "6. 참고문헌", "부록. 연구 상태와 다음 단계",
    ]
    for entry in entries:
        p = doc.add_paragraph()
        set_font(p.add_run(entry), 10.5)
        p.paragraph_format.space_after = Pt(10)
    doc.add_page_break()


def build(template: Path, markdown: Path, output: Path) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template, output)
    doc = Document(output)
    clear_body(doc)
    configure_styles(doc)
    section = doc.sections[0]
    section.top_margin, section.bottom_margin = Cm(2.2), Cm(2.0)
    section.left_margin, section.right_margin = Cm(2.2), Cm(2.0)
    section.footer_distance = Cm(0.9)
    section.footer.is_linked_to_previous = False
    set_page_number(section.footer.paragraphs[0])
    title, sections = parse_markdown(markdown)
    add_cover(doc, title)
    doc.add_heading("국문초록", level=1)
    add_text_block(doc, sections["국문초록"])
    doc.add_heading("영문초록", level=1)
    # English text wraps more aggressively under LibreOffice. Keep abstract and
    # keywords together so a keywords-only spill page is never generated.
    add_text_block(doc, sections["Abstract"], body_size=9.8, keyword_size=9.2)
    add_contents(doc)
    order = [("서론", "1. 서론"), ("방법", "2. 연구 방법"), ("결과", "3. 연구 결과"),
             ("고찰", "4. 고찰"), ("결론", "5. 결론"), ("참고문헌 초안", "6. 참고문헌")]
    for key, heading in order:
        doc.add_heading(heading, level=1)
        add_text_block(doc, sections[key])
    doc.add_heading("부록. 연구 상태와 다음 단계", level=1)
    warning = (
        "지도교수 승인, PRESS 35항목, 우선 문헌 118건, 공개 전문 63건, 근거 문단 326건, "
        "원문 정량 통계 124건을 구조화하고 규칙 6건과 독립 시나리오 12건을 확인하였다. "
        "정량 통계 독립 검증·효과합성과 외부 맹검 평가는 미수행이며 본 문서는 제출 전 작업본이다."
    )
    add_text_block(doc, [(9, warning)])
    for text in (
        "완료: PRESS 35항목과 우선 문헌 118건 제목·초록 확인.",
        "완료: 공개 전문 63건 판정과 근거 문단 326건 원문·locator 확인.",
        "완료: KDRI 규칙 6건 released 승격 및 source/locator 연결률 100%.",
    ):
        p = doc.add_paragraph(style="List Bullet")
        set_font(p.add_run(text), 10)
    doc.core_properties.title = title
    doc.core_properties.author = "권혁찬"
    doc.core_properties.subject = "research_v3 evidence-bound thesis draft"
    doc.core_properties.comments = "Generated from traceable research_v3 metrics; not a final submission."
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", type=Path, required=True)
    parser.add_argument("--markdown", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    build(args.template, args.markdown, args.output)


if __name__ == "__main__":
    main()
