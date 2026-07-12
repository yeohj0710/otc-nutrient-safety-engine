#!/usr/bin/env python3
"""Generate the reduced-scope Korean thesis DOCX from machine-readable artifacts."""
from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2] / "research_v2"
OUT = ROOT / "thesis"
TITLE = "고함량 영양성분의 안전성 근거 검토와 개인맞춤형 정보제공 도구의 개발"
TITLE_EN = "Safety Evidence Review of High-Dose Nutrients and Development of a Personalized Information Tool"


def rows(path: Path):
    with path.open(encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tc_pr.append(shd)


def set_cell_text(cell, text: str, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]; r = p.add_run(text); r.bold = bold; r.font.size = Pt(8.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def heading(doc: Document, text: str, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
    return p


def para(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0.55)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(5)
    return p


def page(doc: Document):
    doc.add_page_break()


def evidence_table(doc: Document, data: list[tuple[str, str, str, str]]):
    table = doc.add_table(rows=1, cols=4); table.style = "Table Grid"; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for c, title in zip(table.rows[0].cells, ["연구", "설계", "노출", "주요 결과"]):
        set_cell_text(c, title, True); shade(c, "DCE6F1")
    widths = [Cm(3.0), Cm(2.5), Cm(3.2), Cm(8.0)]
    for row_data in data:
        cells = table.add_row().cells
        for c, value, width in zip(cells, row_data, widths):
            c.width = width; set_cell_text(c, value)
    doc.add_paragraph("")


def main():
    metrics = json.loads((OUT / "metrics_manifest.json").read_text(encoding="utf-8"))["metrics"]
    evidence_map = rows(ROOT / "synthesis" / "abstract_evidence_map.csv")
    evidence = rows(ROOT / "extraction" / "seed_abstract_evidence.csv")
    screening = json.loads((ROOT / "screening" / "computational_screening_summary.json").read_text(encoding="utf-8"))

    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21), Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.0)
    sec.top_margin, sec.bottom_margin = Cm(2.0), Cm(1.8)
    styles = doc.styles
    styles["Normal"].font.name = "Malgun Gothic"; styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    styles["Normal"].font.size = Pt(10)
    for name in ("Title", "Heading 1", "Heading 2"):
        styles[name].font.name = "Malgun Gothic"; styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        styles[name].font.color.rgb = RGBColor(25, 35, 55)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("졸 업 논 문"); r.bold = True; r.font.size = Pt(18)
    doc.add_paragraph("")
    cover = doc.add_table(rows=8, cols=2); cover.style = "Table Grid"; cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover_data = [
        ("연구 주제명", f"국문: {TITLE}\n영문: {TITLE_EN}"),
        ("성 명", "권혁찬"), ("학 번", "2021194024"),
        ("실습기간", "2026년 3월 3일 ~ 2026년 6월 19일"),
        ("실습장소", "연세대학교 약학대학"), ("논문양식", "종설논문"),
        ("담당교수", "장민정 교수님  (서명)"), ("제출일", "2026년 6월"),
    ]
    for row, (label, value) in zip(cover.rows, cover_data):
        set_cell_text(row.cells[0], label, True); set_cell_text(row.cells[1], value)
        shade(row.cells[0], "E7E6E6")
    doc.add_paragraph("")
    p = doc.add_paragraph("연세대학교 약학대학 전공심화실습(1)"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True; p.runs[0].font.size = Pt(12)

    page(doc); heading(doc, "국문초록")
    para(doc, f"본 연구는 일반의약품 및 건강기능식품으로 접하기 쉬운 비타민 D, 비타민 B6, 철분, 마그네슘 및 아연을 대상으로 안전성 근거를 검토하고, 복용자의 상태에 따라 확인해야 할 정보를 제시하는 조회 도구를 개발하고자 수행하였다. 문헌검색은 PubMed를 이용하였으며, 각 성분에서 임상적으로 문제가 될 수 있는 이상반응과 위험요인을 중심으로 검색식을 구성하였다.")
    para(doc, f"다섯 성분의 검색 결과는 모두 {metrics['pubmed_occurrences']['value']:,}건이었고, 성분 간 중복 문헌 {metrics['duplicates_removed']['value']:,}건을 제외한 뒤 {metrics['unique_records']['value']:,}건을 분석 자료로 정리하였다. 검색식의 누락 여부를 확인하기 위해 미리 선정한 핵심문헌 22편의 회수 여부를 평가하였다. 최초 검색에서는 15편이 확인되었으나 누락 원인을 분석해 검색어를 보완한 뒤 22편을 모두 회수하였다.")
    para(doc, "핵심문헌의 초록을 검토한 결과, 비타민 D에서는 고칼슘혈증·고칼슘뇨·요로결석, 비타민 B6에서는 감각신경병증, 철분에서는 위장관 이상반응, 마그네슘에서는 설사와 고마그네슘혈증, 아연에서는 구리결핍에 따른 혈액학적·신경학적 이상이 주요 안전성 문제로 확인되었다. 이를 바탕으로 성분별 복용량, 복용기간, 병용제품, 기저질환 및 증상을 확인하는 정보제공 문안을 작성하였다.")
    para(doc, "본 연구는 안전성 문헌을 성분별 상담 질문과 연결했다는 점에서 의의가 있다. 다만 PubMed 이외의 데이터베이스 검색과 독립적인 전문 검토를 시행하지 않았으므로, 개발한 문안은 약사의 판단을 보조하는 연구용 자료로 사용되어야 한다.")
    para(doc, "주요어: 일반의약품, 영양제, 안전성, PubMed, 비타민 D, 비타민 B6, 철분, 마그네슘, 아연")

    heading(doc, "영문초록")
    para(doc, "This study reviewed safety evidence for vitamin D, vitamin B6, oral iron, magnesium, and zinc and developed a personalized information tool. Five PubMed searches identified 16,194 records in total, and 15,890 unique records remained after removal of 304 duplicates. Search terms were revised after an initial check retrieved 15 of 22 prespecified key articles; the revised searches retrieved all 22 articles. Major safety concerns were hypercalcemia and urinary stones for vitamin D, sensory neuropathy for vitamin B6, gastrointestinal adverse events for oral iron, diarrhea and hypermagnesemia for magnesium, and copper-deficiency-related hematologic and neurologic abnormalities for zinc. The resulting information statements prompt users to confirm dose, duration, concurrent products, underlying disease, and relevant symptoms. The tool is intended to support, not replace, professional judgment.")
    para(doc, "Keywords: over-the-counter products, dietary supplements, safety, PubMed, evidence traceability")

    page(doc); heading(doc, "목차")
    for item in ["국문초록", "영문초록", "1. 서론", "2. 연구 방법", "  2.1 연구설계와 범위", "  2.2 검색과 계보", "  2.3 중복 제거와 계산 선별", "  2.4 초록 근거와 규칙", "  2.5 연구질문의 구체화", "  2.6 검색전략 설계", "  2.7 핵심문헌 검증", "  2.8 자료 처리", "  2.9 근거 합성", "  2.10 도구 설계", "  2.11 평가변수", "3. 연구 결과", "4. 고찰", "5. 결론", "6. 참고문헌", "부록"]:
        doc.add_paragraph(item)

    page(doc); heading(doc, "1. 서론")
    para(doc, "비타민과 무기질 보충제는 처방전 없이 구입할 수 있어 비교적 안전한 제품으로 인식되기 쉽다. 그러나 같은 성분을 여러 제품으로 복용하거나 고함량 제품을 장기간 사용하는 경우에는 이상반응이 나타날 수 있다. 특히 신기능 저하, 결석 병력, 빈혈 또는 신경학적 증상이 있는 복용자에서는 제품명만으로 안전성을 판단하기 어렵다.")
    para(doc, "영양성분의 안전성은 성분의 유무보다 실제 섭취량과 기간, 병용제품 및 복용자의 임상적 특성에 의해 달라진다. 비타민 D의 과량 섭취는 칼슘대사 이상과 결석 위험을, 비타민 B6의 장기 복용은 감각신경병증을 일으킬 수 있다. 철분은 위장관 이상반응이 흔하고, 마그네슘은 신기능이 저하된 환자에서 축적될 수 있으며, 아연의 과량 섭취는 구리 흡수를 방해할 수 있다.")
    para(doc, "따라서 영양성분 상담에서는 단순히 복용 가능 여부를 답하기보다 제품별 함량을 합산하고, 복용기간과 위험요인을 확인하며, 판단의 근거가 되는 문헌을 함께 제시할 필요가 있다. 본 연구의 목적은 다섯 영양성분의 주요 안전성 근거를 정리하고, 이를 실제 상담에서 확인할 질문으로 전환한 개인맞춤형 정보제공 도구를 개발하는 것이다.")

    heading(doc, "2. 연구 방법")
    heading(doc, "2.1 연구설계와 범위", 2)
    para(doc, "본 연구는 문헌고찰과 정보제공 도구 개발을 결합한 탐색적 연구로 수행하였다. 연구대상 성분은 비타민 D, 비타민 B6, 철분, 마그네슘 및 아연으로 정하였다. 각 성분에서 용량 또는 장기 복용과 관련성이 알려진 안전성 결과를 별도의 연구주제로 설정하였다.")
    heading(doc, "2.2 검색과 계보", 2)
    para(doc, "각 노드 검색식은 MeSH와 제목·초록 자유어를 조합하고 언어·기간·연구설계 필터를 적용하지 않았다. NCBI E-utilities를 사용해 UID와 XML을 전량 저장했다. 각 원시 파일의 SHA-256, 쿼리 해시, hit·export·import 수를 manifest에 기록하고 세 수가 일치하지 않으면 실행을 실패시켰다.")
    para(doc, "검색식의 민감도를 점검하기 위해 연구 시작 전에 성분별 핵심문헌 22편을 선정하였다. 최초 검색에서 확인되지 않은 7편은 제목, 초록 및 MeSH 용어를 다시 검토하였다. 그 결과 과량복용, 이상반응, 투여 및 일반 성분명을 나타내는 표현이 충분히 포함되지 않은 것으로 확인되어 검색어를 보완한 뒤 다시 검색하였다.")
    heading(doc, "2.3 중복 제거와 계산 선별", 2)
    para(doc, "검색 결과는 PMID를 기준으로 중복을 제거하였다. 제목과 초록에 동물실험, 시험관 연구 또는 소아 연구임이 명확하게 제시된 문헌은 우선순위를 낮추었고, 성인에서 경구로 섭취한 영양성분의 이상반응을 다룬 연구를 우선 검토 대상으로 분류하였다. 정보가 충분하지 않은 문헌은 임의로 제외하지 않았다.")
    heading(doc, "2.4 초록 근거와 규칙", 2)
    para(doc, "핵심문헌에서 연구설계, 대상 성분, 주요 안전성 결과와 근거 문장을 추출하였다. 추출한 내용은 PubMed PMID 및 원문 연결주소와 함께 저장하였다. 이후 성분별로 복용자에게 확인할 항목과 상담 시 설명할 내용을 작성하였다. 전문 확인이 끝나지 않은 내용은 확정적인 인과관계나 치료 권고로 표현하지 않았다.")

    heading(doc, "2.5 연구질문의 구체화 과정", 2)
    para(doc, "초기 연구주제는 고함량 영양성분 전반의 안전성을 다루는 것이었다. 그러나 영양성분마다 위해가 나타나는 기전과 임상결과가 달라 하나의 공통 기준으로 비교하면 실제 상담에 필요한 정보가 소실될 수 있었다. 이에 연구질문을 ‘어떤 성분이 위험한가’가 아니라 ‘각 성분에서 어떤 복용조건과 환자 특성을 확인해야 하는가’로 구체화하였다.")
    para(doc, "성분별 연구질문은 노출, 위험상황 및 임상결과의 세 요소로 구성하였다. 노출에는 성분명, 제형, 용량과 복용기간을 포함하였다. 위험상황에는 칼슘 병용, 신기능 저하, 고령, 장기간 복용 및 여러 제품의 중복 섭취를 포함하였다. 임상결과는 실제 복약상담에서 확인할 수 있는 증상 또는 검사결과로 제한하였다. 이 구조를 사용한 이유는 문헌검색어와 조회 도구의 질문 항목이 동일한 임상 개념을 공유하도록 하기 위해서였다.")
    rationale = doc.add_table(rows=1, cols=4); rationale.style = "Table Grid"; rationale.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c, title in zip(rationale.rows[0].cells, ["성분", "선정한 안전성 문제", "선정 이유", "상담 시 확인 항목"]): set_cell_text(c, title, True); shade(c, "DCE6F1")
    for values in [
        ("비타민 D", "고칼슘혈증·고칼슘뇨·요로결석", "고용량과 칼슘 병용 시 칼슘대사 결과가 달라질 수 있음", "총 섭취량, 칼슘 병용, 결석 병력"),
        ("비타민 B6", "감각신경병증", "장기간 고용량 pyridoxine 노출과 관련된 대표적 위해", "중복제품, 복용기간, 감각·보행 이상"),
        ("철분", "위장관 이상반응", "복약순응도와 치료지속에 직접 영향을 주는 흔한 이상반응", "제제, 원소철 함량, 횟수, 위장관 증상"),
        ("마그네슘", "설사·고마그네슘혈증", "일반적 이상반응과 신기능 저하 환자의 중증 위해를 함께 고려해야 함", "제제, 용량, 기간, 신기능, 의식·혈압 변화"),
        ("아연", "구리결핍·혈액학적/신경학적 이상", "만성 과량복용 후 늦게 발견되고 일부 후유증이 남을 수 있음", "용량, 기간, 빈혈, 감각·보행 이상"),
    ]:
        cells=rationale.add_row().cells
        for c,v in zip(cells,values): set_cell_text(c,v)
    doc.add_paragraph("")

    heading(doc, "2.6 검색전략의 설계와 선택 이유", 2)
    para(doc, "검색전략은 PRISMA-S의 재현 가능한 검색 보고 원칙을 참고하여 데이터베이스, 전체 검색식, 검색일, 검색 건수 및 원시 반입 건수를 기록하도록 설계하였다. 각 검색식은 성분 개념, 경구 복용 또는 보충제 노출 개념, 안전성 결과 개념을 AND로 결합하였다. 각 개념 안에서는 MeSH와 제목·초록 자유어를 OR로 연결하였다. MeSH만 사용하면 최근에 색인되지 않은 문헌이 누락될 수 있고, 자유어만 사용하면 과거 문헌의 표현 차이를 놓칠 수 있어 두 방식을 병용하였다.")
    para(doc, "언어, 출판연도 및 연구설계 필터는 적용하지 않았다. 안전성 문헌은 무작위시험뿐 아니라 증례보고, 코호트 연구와 규제기관 의견서에 분산되어 있고, 고전적 독성 증례가 현재의 상담 기준을 형성한 경우가 있기 때문이다. 검색 결과가 많다는 이유로 관련도 상위 일부만 반입하면 검색엔진 순위에 따른 선택편향이 생길 수 있으므로 모든 PMID와 원문 XML을 반입하였다.")
    para(doc, "PRESS 지침에서 제시하는 연구질문의 검색개념 변환, Boolean 연산자, 자유어, 주제명, 제한조건 및 데이터베이스별 변환 항목을 자체 점검에 사용하였다. 다만 정보전문가의 독립 PRESS 검토를 받은 것은 아니므로 이를 완료된 동료검토로 기술하지 않았다.")

    heading(doc, "2.7 핵심문헌을 이용한 검색식 검증", 2)
    para(doc, "검색 건수만으로 검색식의 민감도를 판단할 수 없기 때문에 성분별로 임상적 관련성이 높은 핵심문헌을 미리 선정하였다. 핵심문헌은 직접적인 성분 노출과 안전성 결과를 다루고 PMID로 확인 가능한 연구로 한정하였다. 비타민 D 5편, 비타민 B6 4편, 철분 4편, 마그네슘 5편, 아연 4편으로 총 22편이었다.")
    para(doc, "검색식 검증의 1차 지표는 핵심문헌 회수율로 정의하였다. 분자는 해당 검색식으로 회수된 핵심문헌 수, 분모는 사전 선정 핵심문헌 수로 계산하였다. 최초 검색에서 누락된 문헌은 단순히 목록에 추가하지 않고 어느 검색개념에서 실패했는지 확인하였다. 이후 누락 원인을 설명할 수 있는 동의어만 추가하고 검색식을 새 버전으로 저장한 뒤 전체 검색을 다시 수행하였다. 이 절차는 결과를 알고 검색어를 무제한으로 늘리는 것을 피하면서도 명백한 표현 누락을 교정하기 위해 선택하였다.")

    page(doc); heading(doc, "2.8 자료 처리와 문헌 우선순위 분류", 2)
    para(doc, "각 검색 결과는 검색 실행 단위를 보존한 상태로 저장한 뒤 PMID를 기준으로 성분 간 중복을 확인하였다. 같은 문헌이 둘 이상의 성분 검색에서 확인된 경우 문헌 자체는 하나로 유지하되 해당 성분 정보를 모두 연결하였다. 검색 건수, 반출 PMID 수 및 정규화된 문헌 수가 일치하지 않으면 분석을 중단하도록 하여 자료 누락을 확인하였다.")
    para(doc, "15,000편이 넘는 문헌을 동일한 깊이로 검토하기 어려워 제목과 초록을 이용한 우선순위 분류를 적용하였다. 이 단계의 목적은 최종 포함 여부를 자동으로 결정하는 것이 아니라 사람이 먼저 읽을 문헌을 정하는 것이었다. 성인, 경구 노출, 용량 또는 기간, 이상반응과 사람 연구설계가 함께 나타난 문헌은 우선 검토 대상으로 분류하였다. 동물·시험관·소아 연구임이 명확한 경우에는 연구범위 외 후보로 분류하였고, 정보가 부족한 경우에는 제외하지 않고 추가 확인 대상으로 남겼다.")
    para(doc, "자동화 분류 결과를 최종 문헌선정 결과로 사용하지 않은 이유는 초록에 노출 형태나 안전성 결과가 생략될 수 있기 때문이다. 따라서 결과표의 ‘우선 검토 문헌’은 포함 연구 수가 아니라 검토 순서를 나타내며, 전문 확인을 마친 최종 근거집합과 구분하였다.")

    heading(doc, "2.9 자료추출과 근거 합성", 2)
    para(doc, "핵심문헌에서는 저자, 연도, 연구설계, 대상자, 성분과 제형, 투여량, 복용기간, 비교군, 안전성 결과, 효과크기 및 95% 신뢰구간을 추출하도록 항목을 정하였다. 초록에서 확인되지 않는 값은 추정하지 않고 미확인으로 기록하였다. 같은 결과라도 연구마다 정의가 다른 경우에는 원문의 정의를 유지하였다.")
    para(doc, "정량적 메타분석은 시행하지 않았다. 성분별 문헌의 설계가 무작위시험, 코호트, 증례군, 증례보고와 과학적 의견서로 이질적이었고, 노출량과 결과 정의도 달랐기 때문이다. 서로 다른 결과를 하나의 효과크기로 결합하면 임상적 의미가 왜곡될 수 있어 SWiM 지침의 원칙에 따라 연구를 성분과 안전성 결과별로 묶고, 각 연구의 방향과 구체적 수치를 표로 제시하는 서술적 합성을 선택하였다.")

    heading(doc, "2.10 정보제공 도구의 설계", 2)
    para(doc, "정보제공 도구는 위험도를 하나의 점수로 계산하는 방식 대신 단계별 확인 질문을 제시하도록 설계하였다. 첫 단계에서는 성분과 제품별 함량을 확인하고, 두 번째 단계에서는 하루 총량과 복용기간을 계산하며, 세 번째 단계에서는 기저질환과 병용제품을 확인하고, 마지막 단계에서는 현재 증상과 의료진 상담 필요성을 제시한다.")
    para(doc, "단일 점수를 사용하지 않은 이유는 같은 섭취량이라도 신기능, 결석 병력 또는 병용제품에 따라 의미가 달라지기 때문이다. 또한 문헌에서 직접 확인되지 않은 용량 기준을 새로 만들지 않고, 각 문안에 근거가 된 PMID를 연결하였다. 결과가 불확실한 경우에는 ‘안전하다’ 또는 ‘위험하다’고 단정하지 않고 추가로 확인해야 할 정보와 상담 필요성을 제시하였다.")
    flow = doc.add_table(rows=1, cols=5); flow.style="Table Grid"; flow.alignment=WD_TABLE_ALIGNMENT.CENTER
    for c,v in zip(flow.rows[0].cells,["① 제품 확인", "② 총량·기간", "③ 위험요인", "④ 증상", "⑤ 근거·행동"]): set_cell_text(c,v,True); shade(c,"DCE6F1")
    row=flow.add_row().cells
    for c,v in zip(row,["성분·제형·표시량", "제품 간 합산", "신기능·병력·병용", "이상반응 징후", "PMID와 상담 권고"]): set_cell_text(c,v)
    doc.add_paragraph("")

    heading(doc, "2.11 평가변수와 해석 원칙", 2)
    para(doc, "본 연구의 기술적 평가변수는 전체 검색 건수, 중복 제거 후 고유 문헌 수, 핵심문헌 회수율, 검색 건수와 반입 건수의 일치 여부 및 근거문헌 연결 여부였다. 핵심문헌 회수율은 검색식 검증 지표이며 자동분류의 민감도로 해석하지 않았다. 자동분류 성능을 평가하려면 독립된 사람이 모든 문헌의 정답을 판정한 자료가 필요하기 때문이다.")
    para(doc, "임상적 민감도, 특이도, 양성예측도, F1 점수 및 치명적 위해 누락률은 독립 사례와 전문가 정답이 없어 계산하지 않았다. 평가되지 않은 지표를 0 또는 100%로 표시하지 않았으며, 개발된 문안도 임상 검증을 마친 규칙으로 표현하지 않았다. 사람 대상 자료를 수집하지 않았고 공개 문헌만 사용하였으므로 본 단계에서는 환자 개인정보를 처리하지 않았다.")

    heading(doc, "3. 연구 결과")
    heading(doc, "3.1 검색과 선별", 2)
    table = doc.add_table(rows=1, cols=5); table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.style = "Table Grid"
    for c, t in zip(table.rows[0].cells, ["성분", "검색 문헌 수", "핵심문헌", "우선 검토 문헌", "전문 검토"]): set_cell_text(c, t, True); shade(c, "DCE6F1")
    node_counts = screening["node_membership_counts"]
    node_labels = {"K1": "비타민 D", "K2": "비타민 B6", "K3": "철분", "K4": "마그네슘", "K5": "아연"}
    design_labels = {"clinical_trial": "임상시험", "randomized_controlled_trial": "무작위배정시험", "case_series": "증례군", "scientific_opinion": "과학적 의견서", "systematic_review": "체계적 문헌고찰", "systematic_review_meta_analysis": "체계적 문헌고찰·메타분석", "cohort": "코호트 연구", "perspective_with_structured_review": "서술적 고찰", "case_report": "증례보고", "review": "종설"}
    for map_row in evidence_map:
        node = map_row["clinical_node_id"]
        cells = table.add_row().cells
        vals = [node_labels[node], f"{node_counts[node]:,}", map_row["prespecified_seed_count"], str(json.loads((ROOT / "extraction" / "abstract_evidence_shortlist_summary.json").read_text())["node_counts"][node]), "미시행"]
        for c, v in zip(cells, vals): set_cell_text(c, v)
    para(doc, f"다섯 성분의 검색 건수는 모두 {screening['identified_occurrences']:,}건이었다. 성분 간 PMID 중복 {screening['duplicates_removed']:,}건을 제외한 고유 문헌은 {screening['unique_records']:,}건이었다. 제목과 초록을 이용한 우선순위 분류에서 {screening['proposal_counts']['priority_include_candidate']:,}건을 우선 검토 대상으로, {screening['proposal_counts']['retain_uncertain']:,}건을 추가 확인 대상으로 분류하였다. 연구범위에서 명백히 벗어난 것으로 판단된 문헌은 {screening['proposal_counts']['explicit_exclude_candidate']:,}건이었다.")
    heading(doc, "3.2 핵심문헌 회수율", 2)
    para(doc, "최초 검색에서는 핵심문헌 22편 중 15편이 확인되어 회수율은 68.2%였다. 비타민 B6와 철분에서 각각 1편, 마그네슘에서 3편, 아연에서 2편이 누락되었다. 누락 문헌에 사용된 표현을 반영해 검색식을 수정한 결과 22편이 모두 검색되어 최종 회수율은 100%였다.")
    page(doc); heading(doc, "3.3 성분별 안전성 근거", 2)
    table = doc.add_table(rows=1, cols=5); table.style = "Table Grid"; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c, t in zip(table.rows[0].cells, ["성분", "확인 문헌", "주요 연구설계", "전문 확인", "근거 정리 방식"]): set_cell_text(c, t, True); shade(c, "DCE6F1")
    for row in evidence_map:
        cells = table.add_row().cells; designs = ", ".join(f"{design_labels.get(k, k)} {v}편" for k,v in json.loads(row["design_counts_json"]).items())
        vals = [node_labels[row["clinical_node_id"]], row["abstract_verified_count"], designs, "미시행", "서술적 정리"]
        for c,v in zip(cells, vals): set_cell_text(c,v)
    para(doc, "K1은 임상시험 중심, K2는 사례군·검토·과학적 의견, K3는 무작위시험과 체계적 문헌고찰, K4는 사례군·코호트·관점 논문, K5는 사례보고와 검토로 구성됐다. 설계와 결과 정의가 달라 메타분석은 수행하지 않았다.")

    heading(doc, "3.3.1 비타민 D", 2)
    para(doc, "비타민 D 관련 핵심문헌은 무작위배정시험 4편과 임상시험 1편이었다. 연구들은 고용량 비타민 D가 결석이나 고칼슘혈증을 일관되게 증가시키는지, 칼슘을 함께 복용할 때 고칼슘뇨가 증가하는지를 평가하였다. 연구대상과 투여법이 서로 달라 결과를 단일 위험도로 합치지는 않았다.")
    evidence_table(doc, [
        ("Malihi 등, 2019\nPMID 31005969", "무작위배정시험", "비타민 D3 100,000 IU/월; 중앙 추적기간 3.3년", "결석 158건(비타민 D 76건, 위약 82건). 첫 결석 발생 HR 0.90(95% CI 0.66-1.23). 비타민 D군 고칼슘혈증은 없었다."),
        ("Billington 등, 2018\nPMID 30180273", "무작위배정시험", "고용량 비타민 D와 칼슘 병용", "고용량군 48명 중 19명에서 고칼슘뇨가 발생하였다. 고칼슘뇨 OR 3.6(95% CI 1.39-9.3); 고칼슘혈증 차이는 없었다."),
        ("Ferroni 등, 2017\nPMID 27765695", "무작위배정시험", "1,000 IU/일 대 50,000 IU/주", "두 군 모두 24시간 요중 칼슘과 calcium oxalate/phosphate 과포화도에 유의한 변화가 없었다."),
        ("Leaf 등, 2012\nPMID 22422535", "임상시험", "결석 환자의 비타민 D 보충", "평균 요중 칼슘은 257에서 255 mg/day로 변화하지 않았다(p=0.91). 일부 환자에서는 증가해 보충 후 모니터링 필요성을 제시하였다."),
        ("Gallagher 등, 2014\nPMID 24937025", "무작위배정시험", "칼슘 1,200 mg/일과 비타민 D", "고칼슘뇨와 고칼슘혈증이 관찰되었으나 비타민 D 용량과 뚜렷한 관계는 확인되지 않았다."),
    ])
    para(doc, "비타민 D 결과는 고용량 투여 자체가 모든 대상자에서 결석을 증가시키는 것은 아니라는 점을 보여준다. 그러나 칼슘 병용군에서 고칼슘뇨가 빈번했고 결석 환자 일부에서 요중 칼슘이 증가했으므로, 상담 시 총 비타민 D 섭취량뿐 아니라 칼슘 병용, 결석 병력 및 검사 결과를 함께 확인해야 한다.")

    heading(doc, "3.3.2 비타민 B6", 2)
    para(doc, "비타민 B6의 핵심 안전성 결과는 장기간 고용량 pyridoxine 섭취와 감각신경병증의 관련성이었다. 고전적 증례군은 과량복용 후 운동실조와 중증 감각기능 장애를 보인 성인 7명을 보고했으며, 복용 중단 후 호전되는 경과를 기술하였다.")
    evidence_table(doc, [
        ("Schaumburg 등, 1983\nPMID 6308447", "증례군", "고용량 pyridoxine의 일일 복용", "성인 7명에서 운동실조와 중증 감각신경계 장애가 발생했으며 복용 중단 후 호전되었다."),
        ("Muhamad 등, 2023\nPMID 37447150", "체계적 문헌고찰", "비타민 B6와 말초신경병증", "고용량 B6 노출 환자에서 감각성 신경병증이 보고되었고 중단 후 주관적 증상 호전이 관찰되었다. 허용량 범위의 위해 근거는 제한적이었다."),
        ("EFSA Panel, 2023\nPMID 37207271", "과학적 의견서", "성인 비타민 B6 섭취", "말초신경병증을 결정적 결과로 사용해 성인 상한섭취량을 12 mg/day로 설정하였다."),
        ("Stein 등, 2021\nPMID 33619867", "체계적 문헌고찰·메타분석", "B군 비타민과 말초신경병증", "B군 비타민 치료의 증상 개선은 연구 간 이질성이 컸고, B6 단독의 치료효과를 확정하기 어려웠다."),
    ])
    para(doc, "비타민 B6는 결핍 치료와 과량복용 위해를 구분해야 한다. 제품 한 개의 표시량만 확인하면 복합비타민, 에너지 제품 및 신경비타민제의 중복 섭취를 놓칠 수 있다. 따라서 하루 총량과 복용기간을 합산하고 저림, 감각저하, 균형장애 또는 보행 이상이 있는지 확인하도록 도구를 구성하였다.")

    page(doc); heading(doc, "3.3.3 경구 철분", 2)
    para(doc, "철분 문헌에서는 위장관 이상반응의 빈도와 제제·복용 일정에 따른 차이가 주요 결과였다. 황산철 메타분석은 위약 또는 정맥 철분보다 위장관 부작용이 증가함을 보였고, 여러 경구 제제를 비교한 종설에서는 제제별 이상반응 빈도에 큰 차이가 보고되었다.")
    evidence_table(doc, [
        ("Tolkien 등, 2015\nPMID 25700159", "체계적 문헌고찰·메타분석", "경구 황산철", "위약 대비 위장관 이상반응 OR 2.32(95% CI 1.74-3.08). 임신부 하위군 OR 3.33(95% CI 1.19-9.28). 용량과 OR의 유의한 관련성은 없었다."),
        ("Cancelo-Hidalgo 등, 2013\nPMID 23252877", "체계적 문헌고찰", "경구 철분 제제 111개 연구, 10,695명", "전체 이상반응은 제제에 따라 4.1-47.0%, 위장관 이상반응은 3.7-43.4%였다."),
        ("Pereira 등, 2014\nPMID 24899360", "무작위배정시험", "황산철 65 mg elemental iron, 1일 2회", "첫 주 증상 보고자의 75%가 황산철군이었다. 증상 수는 황산철군 6.7±1.7, 위약군 1.2±0.5(p=0.01)였다."),
        ("Kampuang 등, 2023\nPMID 37010569", "무작위배정시험", "ferrous fumarate 주 3회 대 1일 3회", "64명에서 12주 혈액학적 반응은 유사했으며 주 3회군에서 위장관 이상반응과 비용이 적었다."),
    ])
    para(doc, "이 결과는 철분 상담에서 ‘복용 가능 여부’보다 제제명, 1회 원소철 함량, 하루 복용 횟수 및 증상을 확인하는 것이 중요함을 보여준다. 오심, 구토, 속쓰림, 복통, 설사 및 변비가 지속되면 임의 중단만 권하기보다 복용 일정과 제제 변경 가능성을 전문가와 상의하도록 안내해야 한다.")

    page(doc); heading(doc, "3.3.4 마그네슘", 2)
    para(doc, "마그네슘은 보충제의 설사와 magnesium oxide 복용자의 고마그네슘혈증이라는 두 가지 안전성 문제가 확인되었다. 특히 고령, 신기능 저하, 고용량 및 장기 복용이 함께 존재할 때 중증 위험이 증가하였다.")
    evidence_table(doc, [
        ("Wakai 등, 2019\nPMID 30805197", "후향적 코호트", "MgO 처방 환자 320명", "75명(23%)에서 고마그네슘혈증. eGFR ≤55.4(OR 3.105), BUN ≥22.4(OR 3.490), MgO ≥1,650 mg/day(OR 1.914), ≥36일 복용(OR 2.198)이 독립 위험요인이었다."),
        ("Yamaguchi 등, 2019\nPMID 30136128", "증례군", "MgO 복용 고령자", "65세 이상 신기능 저하 환자 4명에서 증상성 중증 고마그네슘혈증이 발생했고 1명은 사망하였다."),
        ("Nishikawa 등, 2018\nPMID 29988705", "증례군", "응급 혈액투석 환자 15명", "모두 MgO를 복용했고 14명이 65세 이상이었다. 투석 전 Mg 중앙값 6.0 mg/dL; 9명에서 급성신손상이 동반되었다."),
        ("Costello 등, 2023\nPMID 37487817", "서술적 고찰", "보충 마그네슘 128-1,200 mg/day", "다수 중재연구에서 대조군 대비 설사 차이가 없었으나 일부 연구에서는 설사로 중도탈락이 보고되었다."),
        ("Kawano 등, 2024\nPMID 39218658", "코호트 연구", "MgO 입원환자 256명", "BUN, eGFR, MgO 용량과 낮은 BMI가 고마그네슘혈증 위험과 관련되어 혈중 Mg 모니터링 필요성을 제시하였다."),
    ])
    para(doc, "마그네슘 도구에는 제품과 1일 용량 외에도 신기능과 복용기간을 필수 항목으로 포함하였다. 의식저하, 저혈압, 서맥 또는 호흡곤란은 단순한 위장관 부작용과 구분해야 하며, 권장량 이내 복용이라도 고령자나 급성신손상 환자에서는 혈중 마그네슘 확인이 필요할 수 있다.")

    page(doc); heading(doc, "3.3.5 아연", 2)
    para(doc, "아연의 주요 장기 위해는 구리 흡수 저하에 따른 빈혈, 호중구감소증 및 신경학적 이상이었다. 최근 증례들은 일반의약품 또는 보충제 형태의 아연을 장기간 복용한 고령자에서 중증 구리결핍이 뒤늦게 발견될 수 있음을 보여준다.")
    evidence_table(doc, [
        ("Fosmire, 1990\nPMID 2407097", "종설", "과량 아연 섭취", "만성 과량 섭취에서 구리결핍, 빈혈 및 면역·지질대사 변화 가능성을 정리하였다."),
        ("Wahab 등, 2023\nPMID 37640593", "증례보고", "아연 보충제", "아연 보충제 사용으로 인한 중증 구리결핍 빈혈 사례를 보고하였다."),
        ("Gupta 등, 2023\nPMID 37736439", "증례보고", "76세 여성의 아연 보충", "저구리혈증, 중증 범혈구감소증, 빈혈 및 보행장애가 발생하였다. 아연 중단과 구리 투여 후 혈액학적 이상은 회복되었으나 신경학적 장애는 남았다."),
        ("Hooten 등, 2024\nPMID 38846187", "증례보고", "OTC 비타민·무기질 제품", "74세 여성에서 높은 아연, 낮은 구리와 ceruloplasmin, 환상철적모구 및 호중구감소증이 확인되었다. 제품 중단과 구리 투여 후 호전되었다."),
    ])
    para(doc, "아연 관련 증례는 혈액학적 이상이 회복되어도 신경학적 후유증이 남을 수 있음을 보여준다. 따라서 장기 복용자에게서는 아연 용량과 기간을 확인하고, 설명되지 않는 빈혈·백혈구감소·감각 이상 또는 보행장애가 있을 때 혈중 아연과 구리 검사를 의료진과 상의하도록 안내하였다.")
    heading(doc, "3.4 정보제공 도구의 구성", 2)
    para(doc, "조회 도구는 성분별로 복용량, 복용기간, 병용제품, 기저질환 및 현재 증상을 확인하도록 구성하였다. 비타민 D에서는 칼슘 병용과 결석 병력, 비타민 B6에서는 여러 제품의 중복 섭취와 감각 이상, 철분에서는 제제와 복용 횟수 및 위장관 증상, 마그네슘에서는 신기능과 설사·저혈압·의식 변화, 아연에서는 장기 복용과 빈혈·보행 이상을 주요 확인 항목으로 제시하였다. 모든 문안에는 근거가 된 PubMed 문헌을 연결하였다.")

    page(doc); heading(doc, "4. 고찰")
    para(doc, "본 연구에서는 고함량 영양성분의 안전성 문제를 성분별로 구분하고, 문헌에서 확인된 위해를 실제 상담 질문으로 연결하였다. 검색식 보완 전에는 오래된 증례보고나 최근에 색인된 문헌이 누락되었으나, 누락 문헌의 표현을 직접 확인하여 검색어를 수정함으로써 사전에 선정한 핵심문헌을 모두 회수할 수 있었다.")
    para(doc, "마그네슘과 아연 검색에서는 구체적인 제형명뿐 아니라 일반 성분명을 포함한 뒤 검색량이 증가하였다. 이는 관련 문헌을 놓칠 가능성을 줄이는 대신 검토해야 할 비관련 문헌도 늘어날 수 있음을 의미한다. 본 연구에서는 검색 결과의 일부만 선택하지 않고 전체 결과를 보존한 뒤, 연구대상과 노출 및 안전성 결과가 명확한 문헌부터 검토하였다.")
    para(doc, "초록 근거 지도는 각 노드에 안전성 신호가 존재함을 보여 주지만, 발생률이나 인과 크기를 확정하지 않는다. 특히 사례보고가 많은 K5는 드문 위해를 발견하는 데 유용하지만 분모가 없어 위험률을 제시할 수 없다. K1과 K3의 임상시험도 용량, 대상자, 추적기간과 결과 정의가 달라 초록만으로 통합할 수 없다.")
    para(doc, "조회 도구의 장점은 안전성 문안을 근거문헌과 직접 연결하고, 새로운 문헌이 추가될 때 출처를 다시 확인할 수 있다는 점이다. 그러나 출처가 연결되어 있다는 사실만으로 문안의 임상적 타당성이 보장되는 것은 아니다. 실제 상담에 적용하려면 전문 검토와 약사 평가를 거쳐 표현의 정확성과 위해 누락 여부를 확인해야 한다.")
    heading(doc, "4.1 한계", 2)
    para(doc, "첫째, PubMed 이외의 Embase, CENTRAL, Scopus 등은 검색하지 않았다. 둘째, 사람의 이중 선별과 전문 판정이 없다. 셋째, 초록 정보는 용량·기간·분모·결과 정의가 불완전할 수 있다. 넷째, 검색식 검토자가 독립된 정보전문가가 아니다. 다섯째, 성능 평가용 human gold set과 전문가 시나리오가 없어 임상 성능을 주장할 수 없다.")
    heading(doc, "5. 결론")
    para(doc, "본 연구는 비타민 D, 비타민 B6, 철분, 마그네슘 및 아연의 주요 안전성 문제를 문헌에 근거하여 정리하고, 복용량과 기간, 병용제품, 기저질환 및 증상을 확인하는 개인맞춤형 정보제공 도구를 개발하였다. 수정한 검색식은 사전에 선정한 핵심문헌 22편을 모두 회수하였으며, 각 안전성 문안은 근거문헌과 연결되었다. 향후 전문 검토와 독립적인 사례 평가를 통해 문안의 정확성과 상담 활용성을 확인할 필요가 있다.")

    page(doc); heading(doc, "6. 참고문헌")
    for i, row in enumerate(evidence, 1):
        p = doc.add_paragraph(f"{i}. {row['title']} PubMed PMID: {row['pmid']}. {row['source_url']}")
        p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.first_line_indent = Cm(-0.5); p.paragraph_format.space_after = Pt(3)
    for item in [
        "23. Rethlefsen ML, Kirtley S, Waffenschmidt S, et al. PRISMA-S: an extension to the PRISMA Statement for Reporting Literature Searches in Systematic Reviews. Systematic Reviews. 2021;10:39.",
        "24. McGowan J, Sampson M, Salzwedel DM, et al. PRESS Peer Review of Electronic Search Strategies: 2015 Guideline Statement. Journal of Clinical Epidemiology. 2016;75:40-46.",
        "25. Campbell M, McKenzie JE, Sowden A, et al. Synthesis without meta-analysis (SWiM) in systematic reviews: reporting guideline. BMJ. 2020;368:l6890.",
    ]:
        p=doc.add_paragraph(item); p.paragraph_format.left_indent=Cm(0.5); p.paragraph_format.first_line_indent=Cm(-0.5); p.paragraph_format.space_after=Pt(3)

    heading(doc, "부록. 성분별 상담 확인 항목")
    for text_value in ["비타민 D: 총 섭취량, 칼슘 병용, 결석 병력, 고칼슘혈증 관련 증상", "비타민 B6: 복수 제품의 중복 섭취, 복용기간, 저림·감각저하·보행 이상", "철분: 제제명, 1일 복용 횟수, 치료 목적, 오심·변비·설사·복통", "마그네슘: 제제와 용량, 신기능, 설사, 저혈압·서맥·의식 변화", "아연: 용량과 복용기간, 빈혈·호중구감소증, 감각 또는 보행 이상"]:
        doc.add_paragraph(text_value, style="List Bullet")

    for section in doc.sections:
        footer = section.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("- ")
        fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE"); footer._p.append(fld); footer.add_run(" -")
    out = OUT / "권혁찬_졸업논문_전면개작_최종본.docx"
    doc.save(out)
    print(json.dumps({"docx": str(out), "evidence_references": len(evidence)}, ensure_ascii=False))


if __name__ == "__main__":
    main()
