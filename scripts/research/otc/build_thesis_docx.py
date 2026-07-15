from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


REGULAR = "Pretendard"
MEDIUM = "Pretendard Medium"
SEMIBOLD = "Pretendard SemiBold"
EXTRABOLD = "Pretendard ExtraBold"
NAVY = "17233C"
BLUE = "315B7D"
MUTED = "667085"
LIGHT = "EEF3F7"
BORDER = "C9D3DD"
INK = "111827"
CAUTION = "7A4E00"
CAUTION_FILL = "FFF7E0"

# Academic A4 override of the narrative-proposal preset.
CONTENT_WIDTH_IN = 6.102
CONTENT_WIDTH_DXA = 8787
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_BOTTOM = 90
CELL_MARGIN_START_END = 120


def set_run(run, family=REGULAR, size=10.5, bold=False, color=INK, italic=False):
    run.font.name = family
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), family)


def set_style(style, family, size, color, before, after, line=1.35, bold=False):
    style.font.name = family
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    rfonts = style._element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), family)
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if style.name.startswith("Heading"):
        fmt.keep_with_next = True
        fmt.keep_together = True


def set_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    set_run(run, REGULAR, 9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text, end):
        run._r.append(element)


def set_page_number_start(section, start=1):
    sect_pr = section._sectPr
    page_num = sect_pr.find(qn("w:pgNumType"))
    if page_num is None:
        page_num = OxmlElement("w:pgNumType")
        sect_pr.append(page_num)
    page_num.set(qn("w:start"), str(start))


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (
        ("top", CELL_MARGIN_TOP_BOTTOM),
        ("bottom", CELL_MARGIN_TOP_BOTTOM),
        ("start", CELL_MARGIN_START_END),
        ("end", CELL_MARGIN_START_END),
    ):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), BORDER)
        borders.append(node)


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def calculate_widths(rows):
    columns = len(rows[0])
    maxima = [max(len(row[index]) for row in rows) for index in range(columns)]
    weights = [max(6, min(value, 28)) for value in maxima]
    if columns >= 4:
        weights = [max(5, min(value, 20)) for value in maxima]
    total = sum(weights)
    widths = [max(720, round(CONTENT_WIDTH_DXA * weight / total)) for weight in weights]
    difference = CONTENT_WIDTH_DXA - sum(widths)
    widths[-1] += difference
    if widths[-1] < 720:
        shortage = 720 - widths[-1]
        widths[-1] = 720
        donor = max(range(columns - 1), key=lambda index: widths[index])
        widths[donor] -= shortage
    return widths


def parse_table_line(line):
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_table_separator(line):
    cells = parse_table_line(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def parse_markdown(path: Path):
    lines = path.read_text(encoding="utf-8-sig").splitlines()
    title = lines[0].removeprefix("# ").strip()
    blocks = []
    index = 1
    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            index += 1
            continue
        if line.startswith("| ") and index + 1 < len(lines) and is_table_separator(lines[index + 1].strip()):
            table_rows = [parse_table_line(line)]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_rows.append(parse_table_line(lines[index].strip()))
                index += 1
            blocks.append(("table", table_rows))
            continue
        if line.startswith("## "):
            blocks.append(("h1", line[3:].strip()))
        elif line.startswith("### "):
            blocks.append(("h2", line[4:].strip()))
        elif line.startswith("> "):
            blocks.append(("note", line[2:].strip()))
        elif re.match(r"^[-*] ", line):
            blocks.append(("bullet", line[2:].strip()))
        elif re.match(r"^\d+\. ", line):
            blocks.append(("number", line))
        elif re.match(r"^(표|그림)\s+\d+[.-]", line):
            blocks.append(("caption", line))
        elif line == "<!-- PAGEBREAK -->":
            blocks.append(("pagebreak", ""))
        elif line.startswith("<!--") and line.endswith("-->"):
            pass
        elif not line.startswith("Reference basis for Korean prose rhythm:"):
            blocks.append(("body", line.replace("**", "")))
        index += 1
    return title, blocks


def configure_styles(doc):
    styles = doc.styles
    set_style(styles["Normal"], REGULAR, 10.5, INK, 0, 6, 1.45)
    set_style(styles["Title"], EXTRABOLD, 22, NAVY, 0, 12, 1.05, True)
    set_style(styles["Heading 1"], EXTRABOLD, 16, NAVY, 18, 9, 1.2, True)
    set_style(styles["Heading 2"], SEMIBOLD, 12.5, BLUE, 13, 6, 1.25, True)
    set_style(styles["Heading 3"], MEDIUM, 11, BLUE, 9, 4, 1.25)
    set_style(styles["List Bullet"], REGULAR, 10.3, INK, 0, 4, 1.35)
    set_style(styles["List Number"], REGULAR, 10.3, INK, 0, 4, 1.35)
    for name in ("List Bullet", "List Number"):
        fmt = styles[name].paragraph_format
        fmt.left_indent = Inches(0.38)
        fmt.first_line_indent = Inches(-0.18)


def configure_section(section, body=False, title=""):
    section.page_width, section.page_height = Mm(210), Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(30)
    section.right_margin = Mm(25)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)
    if body:
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header.paragraph_format.space_after = Pt(0)
        set_run(header.add_run(title), MEDIUM, 8.5, color=MUTED)
        set_page_field(section.footer.paragraphs[0])


def add_cover(doc, title, document_label, subtitle):
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    configure_section(section)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(64)
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("연세대학교 약학대학"), SEMIBOLD, 13, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(42)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(document_label), EXTRABOLD, 17, True, NAVY)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    set_run(p.add_run(title), EXTRABOLD, 22, True, NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(76)
    set_run(p.add_run(subtitle), MEDIUM, 11, color=MUTED)

    for value, family, size in (
        ("권혁찬", EXTRABOLD, 13),
        ("학번 2021194024", MEDIUM, 10.5),
        ("2026년 7월", REGULAR, 10.5),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        set_run(p.add_run(value), family, size, family == EXTRABOLD, NAVY if family == EXTRABOLD else MUTED)


def add_toc(doc, title, blocks):
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section(section, body=True, title=title)
    set_page_number_start(section, 1)
    heading = doc.add_paragraph(style="Heading 1")
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading.paragraph_format.space_before = Pt(0)
    set_run(heading.add_run("목차"), EXTRABOLD, 16, True, NAVY)

    for kind, content in blocks:
        if kind != "h1":
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Mm(2)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(7)
        p.paragraph_format.keep_with_next = True
        set_run(p.add_run(content), SEMIBOLD, 10.5, color=NAVY)


def add_note(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CAUTION_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run("연구 상태  "), SEMIBOLD, 9.3, color=CAUTION)
    set_run(p.add_run(text), REGULAR, 9.3, color=CAUTION)


def add_markdown_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    widths = calculate_widths(rows)
    set_table_geometry(table, widths)
    set_table_borders(table)
    repeat_table_header(table.rows[0])
    font_size = 8.2 if len(rows[0]) >= 5 else 8.7
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                set_cell_shading(cell, LIGHT)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.2
            if row_index == 0 or re.fullmatch(r"[\d./%-]+", value.strip()):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_run(
                paragraph.add_run(value),
                SEMIBOLD if row_index == 0 else REGULAR,
                font_size,
                row_index == 0,
                NAVY if row_index == 0 else INK,
            )
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(3)


def build(markdown: Path, output: Path, document_label="졸업논문", subtitle="국내 실제 허가 제품의 성분·함량·복용 조건과 안전성 규칙"):
    title, blocks = parse_markdown(markdown)
    doc = Document()
    configure_styles(doc)
    add_cover(doc, title, document_label, subtitle)
    add_toc(doc, title, blocks)

    in_abstract = False
    for kind, content in blocks:
        if kind == "h1":
            in_abstract = content in {"국문초록", "Abstract"}
            is_appendix = content.startswith("부록 ")
            heading = doc.add_paragraph() if is_appendix else doc.add_heading(content, level=1)
            heading.paragraph_format.space_before = Pt(0)
            heading.paragraph_format.space_after = Pt(8)
            heading.paragraph_format.keep_with_next = True
            heading.paragraph_format.keep_together = True
            heading.paragraph_format.line_spacing = 1.15
            heading.paragraph_format.page_break_before = True
            if is_appendix:
                set_run(heading.add_run(content[3:]), EXTRABOLD, 16, True, NAVY)
        elif kind == "h2":
            doc.add_heading(content, level=2)
        elif kind == "note":
            add_note(doc, content)
        elif kind == "table":
            add_markdown_table(doc, content)
        elif kind == "caption":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(7)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            set_run(p.add_run(content), SEMIBOLD, 9.2, True, NAVY)
        elif kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            set_run(p.add_run(content), REGULAR, 10.3)
        elif kind == "number":
            # Preserve the source number. Word's automatic list style otherwise
            # continues numbering across unrelated sections such as references.
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Mm(7)
            p.paragraph_format.first_line_indent = Mm(-5)
            set_run(p.add_run(content), REGULAR, 10.3)
        elif kind == "pagebreak":
            doc.add_page_break()
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if "`" in content else WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Inches(0 if in_abstract else 0.24)
            p.paragraph_format.keep_together = False
            if in_abstract:
                p.paragraph_format.line_spacing = 1.3
                p.paragraph_format.space_after = Pt(5)
            set_run(p.add_run(content), REGULAR, 10 if in_abstract else 10.5)

    doc.core_properties.title = title
    doc.core_properties.author = "권혁찬"
    doc.core_properties.subject = "국내 일반의약품 제품명 중심 안전성 조회 시스템 연구"
    doc.core_properties.keywords = "일반의약품, 제품명 검색, 중복복용, 근거 추적"
    doc.core_properties.comments = "블라인드 독립평가 미완료; 임상 성능 주장 불가."
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--markdown", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--document-label", default="졸업논문")
    parser.add_argument("--subtitle", default="국내 실제 허가 제품의 성분·함량·복용 조건과 안전성 규칙")
    args = parser.parse_args()
    build(args.markdown, args.output, args.document_label, args.subtitle)


if __name__ == "__main__":
    main()
