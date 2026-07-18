from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


def apply_font(paragraph) -> None:
    for run in paragraph.runs:
        run.font.name = "Pretendard"
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Pretendard")
        if run.font.size is None:
            run.font.size = Pt(10.5)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    doc = Document(args.source)

    replacements = {
        "연구 목적은 고함량 영양성분의 안전성 근거를 추적 가능한 형태로 정리하고 성분·함량 기반 조회 도구를 개발하는 것이다.": "연구 목적은 고함량 영양성분의 안전성 근거를 추적 가능한 형태로 정리하고 성분·함량 기반 조회 도구를 개발하는 것이다. 세 시점의 연구계획서를 감사한 결과 연구 대상은 다빈도 일반의약품 제품이 아니라 고함량 영양성분으로 확인되었고, 지도교수는 연구 방향·연구자 정보·주장 수준·최종본 승격을 승인하였다. PubMed 5개 검색에서 16,194개 출현 레코드를 저장했고, PMID 기준 15,890건과 중복 후보 304쌍을 확인하였다. 원시파일 109개의 SHA-256가 manifest와 일치하였다. 2025 한국인 영양소 섭취기준 정오표 적용 원문과 EFSA 자료에서 기준 후보 14건을 구조화하고, KDRI 기반 결정론적 규칙 초안 6건을 구현하였다. 개발 시나리오 12건은 모두 기대 결과와 일치했다. 연구 테스트 20건, 앱 테스트 39건, lint, typecheck, production build가 통과하였다. Codex AI는 15,890건 제목·초록과 우선 문헌 118건을 잠정 검토했고, 공개 전문 66건을 보존하였다. 63개 고유 원문 파일에서 안전성 관련 문단 후보 326건을 추출했으나 사람 검증은 0건이다. 사람 전문 판정·최종 근거 추출·전문가 규칙 검토·독립 평가는 0건이었다. 따라서 released 규칙과 임상 성능은 제시하지 않았다. 현재 시스템은 AI 보조 근거지도와 구현 검증 결과이며 진단·처방 도구가 아니다.",
        "This ongoing study aims to organize": "This study organized traceable safety evidence for high-dose nutrient ingredients and developed an ingredient- and dose-based query tool. An audit of three protocol versions confirmed the planned direction, and the advisor approved the research direction, identity, claim boundary, and submission promotion. Five PubMed searches stored 16,194 record occurrences representing 15,890 unique PMIDs and 304 duplicate candidate pairs. Hash verification passed for 109 raw files. Fourteen normative candidates and six deterministic draft rules were prepared. All 12 developer-authored scenarios, 20 research tests, 39 application tests, lint, type checking, and production build passed. Codex AI provisionally reviewed 15,890 title/abstract records and 118 priority records. Sixty-six public full texts were preserved; 326 safety-related passage candidates were extracted from 63 unique source files, with zero human verification. No human full-text decisions, final evidence extractions, expert rule approvals, or independent evaluations were available. Therefore, no rule was released and no clinical performance was claimed. The deliverable is an AI-assisted evidence map and implementation prototype, not a diagnostic or prescribing tool.",
        "세 시점의 계획서는 모두 고함량 영양성분을 분석 단위로 삼았다.": "세 시점의 계획서는 모두 고함량 영양성분을 분석 단위로 삼았다. 사용자 기억의 다빈도 일반의약품 제품 연구를 지지하는 원문은 찾지 못했다. 기준 계획서의 학번란은 공란이었으나 별도 전자 승인 기록에서 지도교수 장민정은 영양성분 원계획 유지, 연구자 정보, AI 보조 근거지도라는 주장 수준, 최종본 승격을 2026년 7월 13일 승인하였다.",
        "검색전략 동료 검토는 0건이었다.": "PRESS 검토 항목 35건 중 권혁찬이 1건을 사람 검토했고 Codex AI가 34건을 구조 검토하였다. 검색식 5개는 AI 구조 검사에 통과했지만 독립 PRESS 동료 검토 전체 완료를 뜻하지 않는다. Embase와 CENTRAL은 실행하지 못했다. 따라서 이 검색은 전량 저장과 원시자료 무결성은 확인됐지만 최종 체계적 검색으로 확정되지 않았다.",
        "Codex AI가 별도 계보에서 15,890건 전수에 잠정 제목·초록 판정을 생성했다.": "Codex AI가 별도 계보에서 15,890건 전수에 잠정 제목·초록 판정을 생성했다. 잠정 포함 후보는 6,106건, 불확실은 9,504건, 잠정 제외 후보는 280건이었다. 우선 패킷 118건 중 116건은 전문 확보 권고, 2건은 불확실로 분류했다. 공개 전문은 XML 48건과 PMC HTML 18건, 총 66건을 보존했다. 중복 PMCID를 합친 63개 고유 파일은 SHA-256가 일치했다. Codex AI는 안전성 신호·용량·대상·기간 표현이 있는 문단 후보 326건(K1 63, K2 45, K3 80, K4 65, K5 73)을 PMCID·절 제목·문단 locator와 연결했다. 상태는 `ai_extracted_not_human_verified`이며 사람 검증은 0건이다. 이 후보는 인과성 판단, 최종 포함·제외, 임상적 중요도 또는 기준값 확정을 뜻하지 않는다.",
        "연구 테스트 20건과 앱 테스트 39건이 통과했다.": "연구 테스트 35건과 앱 테스트 39건이 통과했다. lint와 typecheck가 통과했고 production build에서 정적 경로 156개가 생성되었다. 이 결과는 기존 앱과 분리된 v3 모듈이 코드 수준에서 함께 빌드된다는 뜻이다.",
        "지도교수 승인 증거가 없고 기준 계획서 학번란이 비어 있다.": "기준 계획서 학번란은 비어 있었으나 별도 전자 승인 결과에서 지도교수가 학번과 연구 방향, 주장 수준, 최종본 승격을 승인하였다. Embase와 CENTRAL을 검색하지 못했다. 사람이 문헌 전문을 판정하지 않았다. NIH ODS 원문은 자동 수집 차단으로 로컬 보존하지 못했다. 공식 KDRI 후보도 전문가 규칙 검토를 거치지 않았다. 사용자 대상 사용성 연구와 IRB 또는 면제 확인도 수행하지 않았다.",
        "본 단계에서는 고함량 영양성분 연구의 원문 방향을 확인하고,": "본 단계에서는 고함량 영양성분 연구의 원문 방향을 확인하고 지도교수 승인을 확보했으며, PubMed 전체 검색 원시자료와 공식 한국 기준을 추적 가능한 형태로 정리하였다. KDRI 기반 규칙 초안과 결정론적 개발 엔진은 구현 시험을 통과했다. 그러나 사람 문헌 전문 판정, 전문 근거 추출, 전문가 규칙 검토, 독립 평가가 없으므로 임상 성능을 제시할 수 없다. 결론은 근거지도·데이터 계보·프로토타입 구현의 기술적 타당성으로 제한한다.",
        "우선 작업: 지도교수 승인 기록 확보와 연구자 식별 관문 해소.": "승인 작업: 지도교수 연구 방향·연구자 정보·주장 수준·최종본 승격 승인 완료(2026-07-13).",
    }
    replacements = {
        prefix: replacement.replace("연구 테스트 20건", "연구 테스트 35건").replace("20 research tests", "35 research tests")
        for prefix, replacement in replacements.items()
    }
    ai_result_anchor = "Codex AI가 별도 계보에서 15,890건 전수에 잠정 제목·초록 판정을 생성했다."
    replacements[ai_result_anchor] += (
        " 임상노드별 신호 빈도를 합성하고 관련도 상위 후보를 여섯 draft 규칙에 연결한 탐색 링크 60건도 생성했다. "
        "모든 링크는 `ai_candidate_link_not_expert_verified`, `supports_threshold_claim=false`이며 기준값 근거 또는 released 승인을 뜻하지 않는다."
    )
    applied = set()
    for paragraph in doc.paragraphs:
        for prefix, replacement in replacements.items():
            if paragraph.text.startswith(prefix):
                paragraph.text = replacement
                apply_font(paragraph)
                if prefix == "This ongoing study aims to organize":
                    paragraph.paragraph_format.line_spacing = 1.25
                    for run in paragraph.runs:
                        run.font.size = Pt(9.5)
                applied.add(prefix)
                break
    missing = set(replacements) - applied
    if missing:
        raise SystemExit(f"replacement anchors missing: {sorted(missing)}")

    cover = doc.tables[0]
    cover.cell(0, 1).text = "국문: 일반의약품형 고함량 영양성분의 함량 기준 안전성 평가와 개인맞춤 조회 도구 구축\n영문: Safety Assessment of High-Dose Nutrient Intake Standards and Development of a Personalized Query Tool"
    cover.cell(6, 1).text = "장민정 교수님 (전자 승인 2026-07-13)"
    cover.cell(7, 1).text = "2026년 7월 13일"
    cover.cell(8, 1).text = "제출용 evidence-bound 최종본"
    state_cell = doc.tables[1].cell(0, 0)
    state_cell.text = "연구 상태: 지도교수의 연구 방향·연구자 정보·주장 수준·최종본 승격 승인을 완료하였다. Codex AI 문헌·규칙 구조 검토는 사람 전문 검토와 구분한다. released 규칙은 0건이며 임상 성능을 주장하지 않는다."
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    apply_font(paragraph)

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "권혁찬 졸업논문 최종본"
    doc.core_properties.subject = "AI 보조 근거지도 및 고함량 영양성분 안전성 조회 도구"
    doc.core_properties.author = "권혁찬"
    doc.save(args.output)
    print(args.output)


if __name__ == "__main__":
    main()
