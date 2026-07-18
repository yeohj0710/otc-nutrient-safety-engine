#!/usr/bin/env python3
"""Insert the deployed prototype methods/results/limitations into the final thesis."""
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[2]
DOCX = ROOT / "research_v2" / "thesis" / "권혁찬_졸업논문_전면개작_최종본.docx"


def add_before(doc, anchor, text, style=None, first_indent=True):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    if first_indent and style is None:
        p.paragraph_format.first_line_indent = Cm(0.55)
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(5)
    anchor._p.addprevious(p._p)
    return p


def add_after(doc, anchor, text):
    p = doc.add_paragraph(text)
    anchor._p.addnext(p._p)
    return p


def main():
    doc = Document(DOCX)
    if any("웹 기반 정보 제공 도구의 구현 및 검증" in p.text for p in doc.paragraphs):
        print("already patched")
        return
    headings = [p for p in doc.paragraphs if p.style.name == "Heading 1"]
    method_anchor = next((p for p in headings if p.text.strip().startswith("3.")), None)
    result_anchor = next((p for p in headings if p.text.strip().startswith("4.")), None)
    discussion_anchor = next((p for p in headings if p.text.strip().startswith("5.")), None)
    if not all((method_anchor, result_anchor, discussion_anchor)):
        raise RuntimeError("본문 장 경계를 찾지 못했습니다.")

    toc = [p for p in doc.paragraphs if p.style.name == "Normal"]
    toc_211 = next((p for p in toc if p.text.strip().startswith("2.11")), None)
    if toc_211 is None:
        raise RuntimeError("정적 목차 경계를 찾지 못했습니다.")
    add_after(doc, toc_211, "2.12 웹 기반 정보 제공 도구의 구현 및 검증")

    add_before(doc, method_anchor, "2.12 웹 기반 정보 제공 도구의 구현 및 검증", "Heading 2", False)
    add_before(doc, method_anchor, "문헌 검색 결과를 일반 사용자가 확인할 수 있도록 웹 기반 연구용 시제품을 구현하였다. 사용자는 성분, 하루 총섭취량, 단위, 복용기간, 병용약 및 질환 정보를 입력한다. 시스템은 입력 성분에 해당하는 규칙만 조회하고, 결과를 ‘먼저 확인할 내용’, ‘함께 확인할 내용’, ‘추가 정보가 필요한 내용’으로 구분한다. 각 결과에는 사용자가 바로 확인할 행동, 수치 기준, 관련 증상과 근거 문헌을 함께 제시하도록 설계하였다. 전문가용 검색 화면을 그대로 노출하지 않고, 초보자도 예시를 선택해 입력과 결과의 관계를 이해할 수 있게 하였다.")
    add_before(doc, method_anchor, "판정 로직은 동일 입력에 동일 결과를 반환하는 결정론적 규칙 엔진으로 구현하였다. 현재 화면에는 17개 연구 성분과 기존 탐색 단계에서 작성된 110개 규칙이 연결되어 있다. 다만 이 규칙은 research_v2의 전문 검토와 출처 위치 확인을 통과한 released 규칙이 아니므로 연구용 legacy 자료로 분리하였다. 최종 released 규칙 수는 0개이며, 이 경계를 화면과 문서에 명시하였다. 자연어 요약 기능은 API 키가 있을 때만 보조적으로 사용할 수 있고, 키가 없거나 호출에 실패하면 검증 가능한 고정 요약문을 출력하도록 하였다. 따라서 생성형 AI의 응답이 판정 수치나 위험 수준을 변경하지 않는다.")
    add_before(doc, method_anchor, "소프트웨어 검증은 단위시험, 연구 산출물 시험, 정적 분석, 형식 검사 및 배포 빌드로 나누어 수행하였다. 규칙 엔진과 화면 동작을 확인하는 Vitest 35건, 검색·중복 제거·근거 계보를 확인하는 Python 시험 51건을 실행하였다. 또한 390 px 너비의 모바일 화면에서 가로 넘침 여부를 확인하고, 마그네슘, 비타민 D, 아연 및 미지원 성분 시나리오로 운영 API의 성분 필터링을 점검하였다.")

    add_before(doc, result_anchor, "3.5 웹 도구 구현 및 검증 결과", "Heading 2", False)
    add_before(doc, result_anchor, "웹 도구는 https://otc-nutrient-safety-engine.vercel.app 에 배포하였다. 배포 빌드는 155개 경로를 생성하였고, Vitest 35건과 Python 연구 시험 51건은 모두 통과하였다. 모바일 390 px 화면에서 가로 스크롤은 발생하지 않았다. 운영 API 시험에서 마그네슘 입력은 마그네슘 결과만, 비타민 D와 티아지드계 이뇨제 입력은 비타민 D 결과만, 아연과 항생제 입력은 아연 결과만 반환하였다. 지원하지 않는 성분은 관련 없는 전체 규칙을 노출하지 않고 결과 0건으로 처리하였다.")
    add_before(doc, result_anchor, "초기 구현에서는 성분을 선택해도 다른 성분의 규칙이 함께 노출되는 오류가 있었다. 이 오류는 사용자가 선택한 성분으로 규칙 집합을 먼저 제한하도록 판정 순서를 수정해 해결하였다. 또한 조회 전부터 결과 카드가 노출되던 화면을 수정하여 실제 조회 후에만 결과를 표시하고, 반복되는 세 개의 설명 상자를 한 문단의 핵심 요약으로 통합하였다. 예시는 실제 지원 성분과 병용 조건을 자동 입력하도록 바꾸어 사용자가 전문 용어를 직접 찾아 입력하지 않아도 주요 흐름을 확인할 수 있게 하였다.")

    add_before(doc, discussion_anchor, "4.2 웹 도구의 해석 범위와 추가 검증 과제", "Heading 2", False)
    add_before(doc, discussion_anchor, "배포된 웹 도구는 문헌 탐색 결과와 인터페이스 설계를 연결한 연구용 시제품이며 진단, 처방 또는 복용 중단을 결정하는 의료기기가 아니다. 화면에 연결된 110개 legacy 규칙은 전문 검토와 locator 확인이 끝난 released 규칙이 아니므로 임상적 정확도를 주장할 수 없다. 현재 연구에서 확인한 시험 통과는 소프트웨어가 정해진 입력을 일관되게 처리하고 연구 수치와 연결된다는 뜻이지, 실제 환자에서 위험을 정확히 예측한다는 뜻은 아니다.")
    add_before(doc, discussion_anchor, "후속 연구에서는 각 규칙의 원문 문장과 위치를 두 명의 검토자가 확인하고, 독립 시나리오에서 민감도와 중대한 위음성을 평가해야 한다. 일반 사용자와 약사를 대상으로 과업 성공률, 결과 이해도, 소요시간 및 오해 발생률을 측정하는 사용성 평가도 필요하다. 이 검증이 끝나기 전에는 사이트의 결과를 전문가 상담을 준비하기 위한 참고 정보로만 해석해야 한다.")

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
