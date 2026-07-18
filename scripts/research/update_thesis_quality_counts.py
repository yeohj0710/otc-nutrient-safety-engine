from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


REPLACEMENTS = {
    "연구 테스트 31건": "연구 테스트 35건",
    "31 research tests": "35 research tests",
    "연구 테스트 33건": "연구 테스트 35건",
    "33 research tests": "35 research tests",
}


def update(source: Path, output: Path) -> int:
    doc = Document(source)
    changed = 0
    containers = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                containers.extend(cell.paragraphs)
    for paragraph in containers:
        for run in paragraph.runs:
            text = run.text
            for old, new in REPLACEMENTS.items():
                if old in text:
                    text = text.replace(old, new)
                    changed += 1
            run.text = text
    if changed != 3:
        raise ValueError(f"expected 3 replacements, found {changed}")
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    return changed


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    print(update(args.source, args.output))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
